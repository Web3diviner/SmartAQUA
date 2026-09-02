"""
Script to generate Chapter 3 Word document for:
Topic: Effect of smart assisted automatic fish feeder on feed management 
       and cost effectiveness in African Catfish farming

Focus: Feed management efficiency and economic analysis
Expanded version with detailed component descriptions matching partner's document length
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_apa_table_borders(table):
    """Apply APA-style borders: top, header-bottom, and table-bottom only"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Remove all borders first
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        if border_name in ['top', 'bottom']:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:color'), '000000')
        else:
            border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # Add bottom border to header row
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)

def add_heading(doc, text, level):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def create_apa_table(doc, headers, data, caption=""):
    """Create an APA-style table with 3 horizontal lines only"""
    if caption:
        cap_para = doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.bold = True
    
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    
    # Add data
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
    
    # Apply APA borders
    set_apa_table_borders(table)
    
    doc.add_paragraph()
    return table

def add_paragraph_text(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para

def main():
    doc = Document()
    
    # Set document margins (1 inch = 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # ========== TITLE ==========
    title = doc.add_heading('CHAPTER 3', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('MATERIALS AND METHODS', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========== 3.1 STUDY AREA ==========
    add_heading(doc, '3.1 Study Area', 2)
    
    doc.add_paragraph(
        "The study was conducted at [LOCATION], [STATE], Nigeria. The experimental site "
        "was selected based on its accessibility, availability of water supply, and proximity "
        "to the research institution. The study area was characterized by tropical climate "
        "conditions with average temperatures ranging from 25°C to 32°C, which falls within "
        "the optimal temperature range for African catfish (Clarias gariepinus) culture "
        "(Kasihmuddin et al., 2021)."
    )
    
    doc.add_paragraph(
        "The experimental facility consisted of two identical concrete tanks, each with a "
        "capacity of 500 liters. The tanks were positioned under a shade structure to minimize "
        "direct sunlight exposure and temperature fluctuations. Water was sourced from a "
        "borehole and stored in an overhead tank before distribution to the experimental units. "
        "The study was conducted over a period of 8 weeks (56 days) from [START DATE] to [END DATE]."
    )
    
    # ========== 3.2 EXPERIMENTAL DESIGN ==========
    add_heading(doc, '3.2 Experimental Design', 2)
    
    doc.add_paragraph(
        "A completely randomized design (CRD) was employed for this study, comprising two "
        "treatment groups: a control group (manual feeding) and a treatment group (smart "
        "automatic feeder). Each treatment was replicated once due to equipment constraints, "
        "with 50 African catfish fingerlings per tank."
    )
    
    # Table: Experimental Design
    create_apa_table(doc, 
        headers=['Parameter', 'Control Group', 'Treatment Group'],
        data=[
            ['Feeding Method', 'Manual', 'Smart Automatic Feeder'],
            ['Feeding Schedule', 'Fixed 3× daily (8:00, 13:00, 18:00)', 'Q10-adjusted dynamic'],
            ['Feed Amount', '3% body weight (fixed)', 'Temperature-adjusted'],
            ['Data Recording', 'Manual logbook', 'Automatic cloud logging'],
            ['Feed Monitoring', 'Visual estimation', 'Load cell measurement'],
            ['Number of Fish', '50', '50'],
            ['Tank Size', '500 liters', '500 liters'],
            ['Duration', '8 weeks (56 days)', '8 weeks (56 days)'],
            ['Feed Type', 'Commercial catfish pellets', 'Commercial catfish pellets'],
            ['Water Exchange', '30% weekly', '30% weekly'],
        ],
        caption='Table 1: Experimental Design Parameters'
    )
    
    # ========== 3.3 EXPERIMENTAL FISH ==========
    add_heading(doc, '3.3 Experimental Fish and Stocking', 2)
    
    doc.add_paragraph(
        "A total of 100 African catfish (Clarias gariepinus) fingerlings were procured from "
        "a reputable hatchery in [LOCATION]. The fingerlings had an average initial weight of "
        "[X] ± [SD] grams and were transported to the experimental site in oxygenated bags. "
        "Upon arrival, the fish were acclimatized for 7 days in a holding tank before being "
        "randomly distributed into the experimental tanks at a stocking density of 50 fish per "
        "500-liter tank (100 fish/m³)."
    )
    
    doc.add_paragraph(
        "During the acclimatization period, fish were fed commercial catfish pellets at 3% "
        "body weight per day to allow adaptation to the experimental conditions. Water quality "
        "parameters were monitored daily to ensure optimal conditions for the fish. Any "
        "mortality during acclimatization was replaced to maintain the stocking density."
    )
    
    # ========== 3.4 SMART AUTOMATIC FISH FEEDER SYSTEM ==========
    add_heading(doc, '3.4 Smart Automatic Fish Feeder System', 2)
    
    doc.add_paragraph(
        "The smart automatic fish feeder system was designed and constructed specifically for "
        "this study. The system integrated Internet of Things (IoT) technology with precision "
        "feeding mechanisms to enable automated, temperature-adjusted feed dispensing with "
        "real-time monitoring capabilities."
    )
    
    add_heading(doc, '3.4.1 System Architecture', 3)
    
    doc.add_paragraph(
        "The smart feeder system comprised three main layers: hardware layer (sensors and "
        "actuators), cloud infrastructure (data processing and storage), and mobile application "
        "(user interface and control). The hardware layer was installed at the fish pond, while "
        "the cloud infrastructure was hosted on Railway platform. Communication between the "
        "hardware and cloud was established via LTE cellular network using MQTT protocol."
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 1: System Architecture - Export from Mermaid Live Editor as PNG]"
    )
    
    add_heading(doc, '3.4.2 Hardware Components', 3)
    
    doc.add_paragraph(
        "The hardware components of the smart feeder system were carefully selected based on "
        "reliability, accuracy, and suitability for outdoor aquaculture environments. Table 2 "
        "presents the detailed specifications of all hardware components used in the system."
    )
    
    # Table: Hardware Components
    create_apa_table(doc,
        headers=['Component', 'Specification', 'Quantity', 'Purpose'],
        data=[
            ['LILYGO T-A7670 R2', 'ESP32 + A7670G LTE', '1', 'Main controller with cellular connectivity'],
            ['ESP32-CAM-MB', 'OV2640 2MP Camera', '1', 'Feeding image capture and verification'],
            ['NEMA 23 Stepper Motor', '2.8A, 1.8°/step, 200 steps/rev', '1', 'Auger drive motor'],
            ['DM542 Stepper Driver', '4.2A, 20-50V DC', '1', 'Motor control with microstepping'],
            ['20mm Wood Drill Auger', 'Stainless steel, 20mm pitch', '1', 'Feed dispensing (25g/rev)'],
            ['HX711 + 20kg Load Cell', '24-bit ADC', '1', 'Feed weight measurement'],
            ['DS18B20', 'Waterproof, 1M cable', '1', 'Water temperature sensing'],
            ['JSN-SR04T', 'Waterproof ultrasonic, 25-400cm', '1', 'Water level detection'],
            ['Solar Panel', '50-100W monocrystalline', '1', 'Primary power source'],
            ['Charge Controller', 'MPPT/PWM 10-20A', '1', 'Battery charging'],
            ['12V Battery', '12-20Ah Lead-acid/LiFePO4', '1', 'Backup power'],
            ['LM2596 DC-DC', 'Adjustable step-down', '2', 'Voltage regulation (5V and 3.3V)'],
        ],
        caption='Table 2: Hardware Components and Specifications'
    )
    

    add_heading(doc, '3.4.3 Main Controller Unit', 3)
    
    doc.add_paragraph(
        "The LILYGO T-A7670 R2 development board served as the main controller for the smart "
        "feeder system. This board integrated an ESP32 microcontroller with an A7670G LTE modem, "
        "enabling both local processing and remote communication capabilities. The ESP32 "
        "microcontroller operated at 240 MHz dual-core processor with 520 KB SRAM and 4 MB flash "
        "memory, providing sufficient computational power for sensor data processing, motor "
        "control, and communication tasks."
    )
    
    doc.add_paragraph(
        "The A7670G LTE modem supported 4G LTE Cat-1 connectivity with fallback to 2G networks, "
        "ensuring reliable communication even in areas with limited network coverage. The modem "
        "was configured to establish MQTT connections with the cloud backend for real-time data "
        "transmission and command reception."
    )
    
    add_heading(doc, '3.4.4 Feed Dispensing Mechanism', 3)
    
    doc.add_paragraph(
        "The feed dispensing mechanism consisted of a NEMA 23 stepper motor coupled to a 20mm "
        "wood drill auger. The stepper motor was controlled by a DM542 driver configured for "
        "8× microstepping, resulting in 1600 effective steps per revolution. The auger was "
        "calibrated to dispense approximately 25 grams of feed per revolution, enabling precise "
        "control of feed quantities."
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 2: Feed Management Data Flow - Export from Mermaid Live Editor as PNG]"
    )
    
    # Table: Motor Control Parameters
    create_apa_table(doc,
        headers=['Parameter', 'Value', 'Unit', 'Description'],
        data=[
            ['Motor Type', 'NEMA 23', '-', 'Bipolar stepper motor'],
            ['Steps per Revolution', '200', 'steps', '1.8 degrees per step'],
            ['Microstepping', '8', 'microsteps', 'Set via DM542 DIP switches'],
            ['Effective Steps per Rev', '1600', 'steps', '200 × 8 microsteps'],
            ['Motor Current', '2.8', 'Amperes', 'Peak current per phase'],
            ['Pulse Width', '5', 'μs', 'Minimum HIGH time for step pulse'],
            ['Step Delay', '1250', 'μs', 'Time between steps at 800 steps/sec'],
            ['Max Speed', '800', 'steps/second', 'Limited for torque at load'],
            ['Motor RPM', '30', 'RPM', '(800 × 60) / 1600 steps'],
            ['Grams per Revolution', '25', 'grams', 'Calibrated feed output'],
            ['Dispensing Rate', '750', 'grams/minute', '30 RPM × 25g per rev'],
        ],
        caption='Table 3: Motor Control Parameters'
    )
    
    add_heading(doc, '3.4.5 Sensor Systems', 3)
    
    doc.add_paragraph(
        "The smart feeder system incorporated multiple sensors for environmental monitoring "
        "and feed management. The DS18B20 waterproof temperature sensor was submerged in the "
        "fish tank to continuously monitor water temperature with an accuracy of ±0.5°C. "
        "Temperature readings were taken every 5 minutes and used to adjust feeding rates "
        "according to the Q10 algorithm."
    )
    
    doc.add_paragraph(
        "The HX711 analog-to-digital converter coupled with a 20kg load cell was used to "
        "measure the weight of feed in the hopper. This enabled accurate tracking of feed "
        "consumption and automatic low-feed alerts when the hopper level dropped below 10%. "
        "The load cell was calibrated using known weights before the experiment commenced."
    )
    
    doc.add_paragraph(
        "The JSN-SR04T waterproof ultrasonic sensor was installed above the tank to monitor "
        "water level. This sensor provided early warning of water level changes that could "
        "affect fish welfare or indicate potential leaks."
    )
    
    # Table: Sensor Specifications
    create_apa_table(doc,
        headers=['Sensor', 'Model', 'Range', 'Accuracy', 'Interface', 'Purpose'],
        data=[
            ['Temperature', 'DS18B20', '-55 to 125°C', '±0.5°C', 'OneWire', 'Water temperature monitoring'],
            ['Weight', 'HX711 + Load Cell', '0-20kg', '±0.1g', 'Digital SPI', 'Feed level measurement'],
            ['Water Level', 'JSN-SR04T', '25-400cm', '±1cm', 'Ultrasonic', 'Water level detection'],
            ['Camera', 'OV2640', '2MP JPEG', 'VGA/SVGA', 'UART', 'Feeding image capture'],
        ],
        caption='Table 4: Sensor Specifications'
    )
    
    add_heading(doc, '3.4.6 Power System', 3)
    
    doc.add_paragraph(
        "The smart feeder system was powered by a solar-battery hybrid system to ensure "
        "continuous operation in remote pond locations. A 50-100W monocrystalline solar panel "
        "was connected to an MPPT/PWM charge controller, which regulated charging of a 12V "
        "12-20Ah battery. Two LM2596 DC-DC step-down converters provided regulated 5V and 3.3V "
        "outputs for the microcontroller and sensors respectively."
    )
    
    doc.add_paragraph(
        "The power system was designed to provide at least 3 days of autonomous operation "
        "without solar charging, ensuring reliability during cloudy weather conditions. Power "
        "consumption was optimized through sleep modes and efficient motor control algorithms."
    )
    
    add_heading(doc, '3.4.7 Wiring Connections', 3)
    
    doc.add_paragraph(
        "The electrical connections between components were carefully designed to ensure "
        "reliable operation and minimize electromagnetic interference. Table 5 presents the "
        "motor driver wiring connections, while Table 6 shows the sensor wiring connections."
    )
    
    # Table: Motor Driver Wiring
    create_apa_table(doc,
        headers=['ESP32 Pin', 'DM542 Terminal', 'Function'],
        data=[
            ['GPIO32', 'PUL+', 'Step pulse signal'],
            ['GPIO33', 'DIR+', 'Direction signal'],
            ['GPIO0', 'ENA+', 'Enable signal'],
            ['GND', 'PUL-, DIR-, ENA-', 'Common ground'],
            ['12-48V DC', 'Power terminals', 'Motor power supply'],
        ],
        caption='Table 5: Motor Driver Wiring Connections'
    )
    
    # Table: Sensor Wiring
    create_apa_table(doc,
        headers=['Sensor', 'ESP32 Pin', 'Function'],
        data=[
            ['HX711 DOUT', 'GPIO39 (VN)', 'Load cell data output'],
            ['HX711 SCK', 'GPIO5', 'Load cell clock'],
            ['DS18B20 Data', 'GPIO23', 'Temperature data (with 4.7kΩ pullup)'],
            ['JSN-SR04T TRIG', 'GPIO17', 'Ultrasonic trigger'],
            ['JSN-SR04T ECHO', 'GPIO34', 'Ultrasonic echo'],
        ],
        caption='Table 6: Sensor Wiring Connections'
    )
    
    add_heading(doc, '3.4.8 DM542 DIP Switch Configuration', 3)
    
    doc.add_paragraph(
        "The DM542 stepper driver was configured using DIP switches to set the microstepping "
        "mode and motor current. For this application, 8× microstepping was selected to provide "
        "smooth motor operation and precise feed dispensing control."
    )
    
    # Table: DIP Switch Settings
    create_apa_table(doc,
        headers=['Setting', 'SW1', 'SW2', 'SW3', 'SW4', 'SW5', 'SW6', 'Result'],
        data=[
            ['Microstepping 8×', 'OFF', 'ON', 'ON', '-', '-', '-', '1600 steps/revolution'],
            ['Motor Current 2.8A', '-', '-', '-', 'ON', 'OFF', 'ON', 'Peak current per phase'],
        ],
        caption='Table 7: DM542 DIP Switch Configuration'
    )
    

    # ========== 3.5 SOFTWARE SYSTEM ==========
    add_heading(doc, '3.5 Software System', 2)
    
    doc.add_paragraph(
        "The software system for the smart feeder comprised three main components: embedded "
        "firmware running on the ESP32 microcontroller, a cloud backend server, and a mobile "
        "application for user interaction. All software components were developed specifically "
        "for this study using modern programming frameworks and best practices."
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 10: Software Architecture - Export from Mermaid Live Editor as PNG]"
    )
    
    # Table: Software Technologies
    create_apa_table(doc,
        headers=['Layer', 'Technology', 'Version', 'Purpose'],
        data=[
            ['Firmware', 'PlatformIO', '6.x', 'ESP32 development environment'],
            ['Firmware', 'Arduino Framework', '2.x', 'Hardware abstraction layer'],
            ['Firmware', 'TinyGSM', '0.11.x', 'LTE modem communication'],
            ['Backend', 'Go', '1.21+', 'Server-side programming'],
            ['Backend', 'Gin', '1.9+', 'HTTP web framework'],
            ['Backend', 'PostgreSQL', '15+', 'Relational database'],
            ['Backend', 'Redis', '7+', 'Caching and real-time data'],
            ['Mobile', 'Flutter', '3.16+', 'Cross-platform UI framework'],
            ['Mobile', 'Dart', '3.2+', 'Programming language'],
            ['Mobile', 'Riverpod', '2.4+', 'State management'],
        ],
        caption='Table 8: Software Technologies'
    )
    
    add_heading(doc, '3.5.1 Q10 Temperature Adjustment Algorithm', 3)
    
    doc.add_paragraph(
        "The Q10 algorithm was implemented to adjust feeding rates based on water temperature. "
        "The Q10 coefficient represents the factor by which metabolic rate increases for every "
        "10°C rise in temperature. For African catfish, a Q10 value of 2.0 was used based on "
        "published literature (Kasihmuddin et al., 2021)."
    )
    
    doc.add_paragraph(
        "The adjusted feeding rate was calculated using the following formula:"
    )
    
    doc.add_paragraph(
        "FR_adj = SFR × Q10^((T - T_ref) / 10) × TIF"
    )
    
    doc.add_paragraph(
        "Where FR_adj is the adjusted feeding rate (g/day), SFR is the standard feeding rate "
        "(g/day), Q10 is the temperature coefficient (2.0), T is the current water temperature "
        "(°C), T_ref is the reference temperature (25°C), and TIF is the thermal inhibition "
        "factor applied when temperature exceeds optimal range."
    )
    
    # Table: Q10 Parameters
    create_apa_table(doc,
        headers=['Parameter', 'Symbol', 'Value', 'Unit', 'Description'],
        data=[
            ['Q10 Coefficient', 'Q10', '2.0', 'dimensionless', 'Metabolic rate temperature sensitivity'],
            ['Reference Temperature', 'T_ref', '25', '°C', 'Standard temperature for calculations'],
            ['Optimal Temperature Range', 'T_opt', '25-30', '°C', 'Best growth temperature for African catfish'],
            ['Critical Low Temperature', 'T_crit_low', '20', '°C', 'Feeding reduced below this temperature'],
            ['Critical High Temperature', 'T_crit_high', '30', '°C', 'Feeding reduced above this temperature'],
            ['Thermal Inhibition Factor', 'TIF', '0.3-0.8', 'dimensionless', 'Penalty during thermal stress'],
            ['Standard Feeding Rate', 'SFR', '3-5', '% BW/day', 'Base feeding rate by body weight'],
        ],
        caption='Table 9: Q10 Algorithm Parameters'
    )
    
    add_heading(doc, '3.5.2 Motor Control Calculations', 3)
    
    doc.add_paragraph(
        "The motor control system calculated the number of steps required to dispense a "
        "specific amount of feed. The calculation involved converting the target feed amount "
        "to motor revolutions, then to step pulses based on the microstepping configuration."
    )
    
    doc.add_paragraph(
        "Total Steps = (Target Grams / Grams per Revolution) × Steps per Revolution × Microsteps"
    )
    
    # Table: Feed Dispensing Calculation Example
    create_apa_table(doc,
        headers=['Step', 'Calculation', 'Value', 'Unit'],
        data=[
            ['Target Feed Amount', 'User input', '100', 'grams'],
            ['Q10 Adjustment Factor', 'Q10^((T-25)/10)', '1.15', 'dimensionless'],
            ['Adjusted Feed Amount', '100 × 1.15', '115', 'grams'],
            ['Revolutions Needed', '115 / 25', '4.6', 'revolutions'],
            ['Steps per Revolution', '200 × 8', '1600', 'steps'],
            ['Total Steps Required', '4.6 × 1600', '7360', 'steps'],
            ['Time per Step', '1 / 800', '1.25', 'milliseconds'],
            ['Total Dispensing Time', '7360 × 1.25', '9.2', 'seconds'],
        ],
        caption='Table 10: Feed Dispensing Calculation Example'
    )
    
    # ========== 3.6 FEEDING MANAGEMENT ==========
    add_heading(doc, '3.6 Feeding Management', 2)
    
    doc.add_paragraph(
        "Commercial floating catfish pellets (Coppens or equivalent) with 42% crude protein "
        "content were used throughout the experiment. The feed was stored in airtight containers "
        "to prevent moisture absorption and maintain nutritional quality."
    )
    
    add_heading(doc, '3.6.1 Control Group (Manual Feeding)', 3)
    
    doc.add_paragraph(
        "Fish in the control group were fed manually three times daily at 08:00, 13:00, and "
        "18:00 hours. The daily feed ration was calculated as 3% of the total body weight of "
        "fish in the tank. Feed amounts were adjusted weekly based on the estimated biomass "
        "from sampling. Feed was weighed using a digital scale (±0.1g accuracy) before each "
        "feeding event and broadcast evenly across the water surface."
    )
    
    doc.add_paragraph(
        "Feeding records were maintained manually in a logbook, documenting the date, time, "
        "feed amount, and any observations about fish feeding behavior. Labor time for each "
        "feeding activity was recorded using a stopwatch to enable comparison with the "
        "automated system."
    )
    
    add_heading(doc, '3.6.2 Treatment Group (Smart Feeder)', 3)
    
    doc.add_paragraph(
        "Fish in the treatment group were fed using the smart automatic feeder system. The "
        "feeding schedule was dynamically adjusted based on water temperature using the Q10 "
        "algorithm. The system automatically calculated the optimal feed amount based on "
        "current temperature and dispensed feed at programmed intervals."
    )
    
    doc.add_paragraph(
        "All feeding events were automatically logged to the cloud database, including "
        "timestamp, target feed amount, actual feed dispensed (measured by load cell), water "
        "temperature, and feeding image captured by the ESP32-CAM. This automated data "
        "collection eliminated manual recording errors and provided precise feed management data."
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 8: Feed Dispensing Mechanism Flowchart - Export from Mermaid Live Editor as PNG]"
    )
    

    # ========== 3.7 DATA COLLECTION ==========
    add_heading(doc, '3.7 Data Collection', 2)
    
    doc.add_paragraph(
        "Data collection was conducted systematically throughout the 8-week experimental "
        "period. Both feed management data and cost data were collected to enable comprehensive "
        "analysis of the smart feeder's effectiveness."
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 11: Data Collection Timeline - Export from Mermaid Live Editor as PNG]"
    )
    
    add_heading(doc, '3.7.1 Feed Management Data', 3)
    
    doc.add_paragraph(
        "Feed management data were collected to evaluate the efficiency and accuracy of the "
        "feeding systems. The following parameters were recorded:"
    )
    
    # Table: Feed Management Parameters
    create_apa_table(doc,
        headers=['Parameter', 'Formula', 'Unit', 'Description'],
        data=[
            ['Total Feed Consumed', 'Direct measurement', 'kg', 'Total feed used during experiment'],
            ['Daily Feed Consumption', 'Total Feed / Days', 'kg/day', 'Average daily feed usage'],
            ['Feed Wastage', 'Feed Dispensed - Feed Consumed', 'kg', 'Uneaten or lost feed'],
            ['Feed Wastage Rate', '(Wastage / Dispensed) × 100', '%', 'Percentage of wasted feed'],
            ['Feed Conversion Ratio', 'Feed Consumed / Weight Gain', 'ratio', 'Feed efficiency measure'],
            ['Feed Accuracy', '(Actual / Target) × 100', '%', 'Dispensing precision'],
            ['Feeding Frequency', 'Count per day', 'times/day', 'Number of feeding events'],
            ['Feed Cost per kg Gain', 'Feed Cost / Weight Gain', '₦/kg', 'Economic feed efficiency'],
        ],
        caption='Table 11: Feed Management Parameters'
    )
    
    add_heading(doc, '3.7.2 Labor and Time Data', 3)
    
    doc.add_paragraph(
        "Labor time was recorded for all feeding-related activities in both treatment groups. "
        "This included feed preparation, actual feeding time, monitoring, and record keeping. "
        "Time was measured using a digital stopwatch and recorded in minutes per activity."
    )
    
    # Table: Labor Requirements
    create_apa_table(doc,
        headers=['Activity', 'Control Group', 'Treatment Group', 'Unit'],
        data=[
            ['Daily Feeding Time', '45-60', '5-10', 'minutes/day'],
            ['Feed Preparation', '15-20', '0', 'minutes/day'],
            ['Monitoring Time', '30-45', '5-10', 'minutes/day'],
            ['Record Keeping', '15-20', '0 (automatic)', 'minutes/day'],
            ['Total Daily Labor', '105-145', '10-20', 'minutes/day'],
            ['Weekly Labor Hours', '12.25-16.92', '1.17-2.33', 'hours/week'],
        ],
        caption='Table 12: Labor and Time Requirements'
    )
    
    add_heading(doc, '3.7.3 Cost Data', 3)
    
    doc.add_paragraph(
        "Cost data were collected throughout the experiment to enable economic analysis. "
        "All costs were recorded in Nigerian Naira (₦) and categorized into fixed costs "
        "(equipment, installation) and variable costs (feed, labor, operating expenses)."
    )
    
    # Table: Cost Data Sources
    create_apa_table(doc,
        headers=['Cost Component', 'Data Source', 'Recording Frequency'],
        data=[
            ['Feed Price', 'Local feed suppliers', 'Per purchase'],
            ['Labor Rate', 'Farm records / Minimum wage', 'Weekly'],
            ['Electricity Cost', 'Utility bills', 'Monthly'],
            ['Equipment Cost', 'Purchase receipts', 'One-time'],
            ['Maintenance Cost', 'Service records', 'As incurred'],
            ['Fish Market Price', 'Local fish markets', 'Weekly'],
            ['Fingerling Cost', 'Hatchery records', 'Initial purchase'],
        ],
        caption='Table 13: Cost Data Sources'
    )
    
    add_heading(doc, '3.7.4 Equipment Cost Breakdown', 3)
    
    doc.add_paragraph(
        "The total equipment cost for the smart feeder system was calculated based on "
        "component prices at the time of purchase. Depreciation was calculated using the "
        "straight-line method based on estimated useful life of each component."
    )
    
    # Table: Equipment Cost Breakdown
    create_apa_table(doc,
        headers=['Component', 'Unit Cost (₦)', 'Quantity', 'Total Cost (₦)', 'Lifespan (Years)'],
        data=[
            ['LILYGO T-A7670 R2', '25,000', '1', '25,000', '5'],
            ['ESP32-CAM-MB', '5,000', '1', '5,000', '5'],
            ['NEMA 23 Stepper Motor', '15,000', '1', '15,000', '5'],
            ['DM542 Stepper Driver', '8,000', '1', '8,000', '5'],
            ['20mm Auger Bit', '3,000', '1', '3,000', '3'],
            ['HX711 + Load Cell', '5,000', '1', '5,000', '5'],
            ['DS18B20 Sensor', '1,500', '1', '1,500', '5'],
            ['JSN-SR04T Sensor', '3,000', '1', '3,000', '5'],
            ['Solar Panel 50W', '20,000', '1', '20,000', '10'],
            ['Charge Controller', '5,000', '1', '5,000', '5'],
            ['12V 20Ah Battery', '15,000', '1', '15,000', '3'],
            ['LM2596 Regulators', '1,000', '2', '2,000', '5'],
            ['Enclosure and Wiring', '10,000', '1', '10,000', '5'],
            ['Total Equipment Cost', '-', '-', '118,500', '-'],
        ],
        caption='Table 14: Equipment Cost Breakdown'
    )
    
    add_heading(doc, '3.7.5 Data Collection Schedule', 3)
    
    # Table: Data Collection Schedule
    create_apa_table(doc,
        headers=['Parameter', 'Frequency', 'Method', 'Responsible'],
        data=[
            ['Feed Dispensed', 'Every feeding', 'Load cell / Manual weighing', 'Automatic / Researcher'],
            ['Water Temperature', 'Every 5 minutes', 'DS18B20 sensor / Thermometer', 'Automatic / Researcher'],
            ['Fish Weight', 'Weekly', 'Digital scale (sample)', 'Researcher'],
            ['Fish Length', 'Weekly', 'Measuring board (sample)', 'Researcher'],
            ['Mortality', 'Daily', 'Visual count', 'Researcher'],
            ['Labor Hours', 'Daily', 'Time sheet', 'Researcher'],
            ['Feed Cost', 'Per purchase', 'Receipt recording', 'Researcher'],
            ['Water Quality', 'Weekly', 'Test kits', 'Researcher'],
            ['Feeding Images', 'Every feeding', 'ESP32-CAM', 'Automatic'],
            ['Market Prices', 'Weekly', 'Market survey', 'Researcher'],
        ],
        caption='Table 15: Data Collection Schedule'
    )
    

    # ========== 3.8 COST EFFECTIVENESS ANALYSIS ==========
    add_heading(doc, '3.8 Cost Effectiveness Analysis', 2)
    
    doc.add_paragraph(
        "Cost effectiveness analysis was conducted to evaluate the economic viability of the "
        "smart automatic feeder compared to manual feeding. The analysis considered both "
        "fixed and variable costs, as well as potential savings and revenue generation."
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 3: Cost Effectiveness Analysis Flowchart - Export from Mermaid Live Editor as PNG]"
    )
    
    add_heading(doc, '3.8.1 Cost Parameters', 3)
    
    doc.add_paragraph(
        "The following cost parameters were calculated for both treatment groups to enable "
        "comprehensive economic comparison:"
    )
    
    # Table: Cost Effectiveness Parameters
    create_apa_table(doc,
        headers=['Parameter', 'Formula', 'Unit', 'Description'],
        data=[
            ['Total Feed Cost', 'Feed Consumed × Feed Price', '₦', 'Cost of feed used'],
            ['Labor Cost', 'Labor Hours × Hourly Rate', '₦', 'Cost of manual labor'],
            ['Equipment Cost', 'Initial Investment', '₦', 'System purchase/construction cost'],
            ['Operating Cost', 'Electricity + Maintenance', '₦', 'Running costs'],
            ['Total Variable Cost (TVC)', 'Feed + Labor + Operating', '₦', 'Total operational costs'],
            ['Total Cost (TC)', 'TVC + Equipment Depreciation', '₦', 'Total costs including depreciation'],
            ['Revenue', 'Fish Sold × Market Price', '₦', 'Income from fish sales'],
            ['Gross Margin', 'Revenue - TVC', '₦', 'Profit before fixed costs'],
            ['Net Profit', 'Revenue - TC', '₦', 'Final profit after all costs'],
            ['Return on Investment (ROI)', '(Net Profit / TC) × 100', '%', 'Investment efficiency'],
            ['Benefit-Cost Ratio (BCR)', 'Revenue / TC', 'ratio', 'Economic viability indicator'],
            ['Payback Period', 'Equipment Cost / Annual Savings', 'months', 'Time to recover investment'],
        ],
        caption='Table 16: Cost Effectiveness Parameters'
    )
    
    add_heading(doc, '3.8.2 Economic Analysis Formulas', 3)
    
    doc.add_paragraph(
        "The following formulas were used for economic analysis:"
    )
    
    doc.add_paragraph(
        "Return on Investment (ROI):"
    )
    doc.add_paragraph(
        "ROI (%) = ((Revenue - Total Cost) / Total Cost) × 100"
    )
    
    doc.add_paragraph(
        "Benefit-Cost Ratio (BCR):"
    )
    doc.add_paragraph(
        "BCR = Total Revenue / Total Cost"
    )
    doc.add_paragraph(
        "Note: BCR > 1.0 indicates economic viability; BCR < 1.0 indicates economic loss"
    )
    
    doc.add_paragraph(
        "Payback Period:"
    )
    doc.add_paragraph(
        "Payback Period (months) = Equipment Cost / Monthly Savings"
    )
    
    doc.add_paragraph(
        "Annual Depreciation (Straight-Line Method):"
    )
    doc.add_paragraph(
        "Annual Depreciation = (Equipment Cost - Salvage Value) / Useful Life (years)"
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 9: Economic Analysis Decision Tree - Export from Mermaid Live Editor as PNG]"
    )
    
    doc.add_paragraph(
        "[INSERT DIAGRAM 13: Cost-Benefit Analysis Process - Export from Mermaid Live Editor as PNG]"
    )
    
    # ========== 3.9 WATER QUALITY MANAGEMENT ==========
    add_heading(doc, '3.9 Water Quality Management', 2)
    
    doc.add_paragraph(
        "Water quality was maintained within optimal ranges for African catfish throughout "
        "the experimental period. Weekly water exchange of 30% was conducted in both tanks "
        "to maintain water quality. The following parameters were monitored:"
    )
    
    doc.add_paragraph(
        "• Temperature: 25-30°C (monitored continuously by DS18B20 in treatment group)"
    )
    doc.add_paragraph(
        "• pH: 6.5-8.5 (measured weekly using pH meter)"
    )
    doc.add_paragraph(
        "• Dissolved Oxygen: >5 mg/L (measured weekly using DO meter)"
    )
    doc.add_paragraph(
        "• Ammonia: <0.02 mg/L (measured weekly using test kit)"
    )
    
    doc.add_paragraph(
        "Aeration was provided using air pumps and air stones to maintain adequate dissolved "
        "oxygen levels. Any deviations from optimal water quality parameters were addressed "
        "immediately through water exchange or treatment."
    )
    
    # ========== 3.10 STATISTICAL ANALYSIS ==========
    add_heading(doc, '3.10 Statistical Analysis', 2)
    
    doc.add_paragraph(
        "Data collected during the experiment were analyzed using descriptive and inferential "
        "statistics. Microsoft Excel was used for data organization and preliminary analysis, "
        "while SPSS (Statistical Package for Social Sciences) version 26 was used for "
        "statistical tests."
    )
    
    # Table: Statistical Methods
    create_apa_table(doc,
        headers=['Analysis', 'Method', 'Software', 'Purpose'],
        data=[
            ['Descriptive Statistics', 'Mean, SD, SE', 'Microsoft Excel / SPSS', 'Summarize data'],
            ['Normality Test', 'Shapiro-Wilk', 'SPSS', 'Verify normal distribution'],
            ['Variance Test', "Levene's Test", 'SPSS', 'Verify homogeneity of variance'],
            ['Inferential Statistics', 'Independent t-Test', 'SPSS', 'Compare group means'],
            ['Economic Analysis', 'Cost-Benefit Analysis', 'Microsoft Excel', 'Evaluate profitability'],
            ['Significance Level', 'α = 0.05', '-', 'Determine statistical significance'],
        ],
        caption='Table 17: Statistical Analysis Methods'
    )
    
    add_heading(doc, '3.10.1 Independent Samples t-Test', 3)
    
    doc.add_paragraph(
        "The independent samples t-test was used to compare means between the control and "
        "treatment groups. The test statistic was calculated using the following formula:"
    )
    
    doc.add_paragraph(
        "t = (X̄₁ - X̄₂) / √(s²p × (1/n₁ + 1/n₂))"
    )
    
    doc.add_paragraph(
        "Where X̄₁ and X̄₂ are the sample means of group 1 and group 2, s²p is the pooled "
        "variance, and n₁ and n₂ are the sample sizes of group 1 and group 2."
    )
    
    doc.add_paragraph(
        "Pooled Variance:"
    )
    doc.add_paragraph(
        "s²p = ((n₁-1)s₁² + (n₂-1)s₂²) / (n₁ + n₂ - 2)"
    )
    
    doc.add_paragraph(
        "Degrees of Freedom:"
    )
    doc.add_paragraph(
        "df = n₁ + n₂ - 2"
    )
    
    doc.add_paragraph(
        "Results were considered statistically significant at p < 0.05. All data were "
        "presented as mean ± standard deviation (SD) unless otherwise stated."
    )
    

    # ========== 3.11 ETHICAL CONSIDERATIONS ==========
    add_heading(doc, '3.11 Ethical Considerations', 2)
    
    doc.add_paragraph(
        "This study was conducted in accordance with ethical guidelines for the use of "
        "animals in research. The experimental protocol was reviewed and approved by "
        "[INSTITUTION NAME] Ethics Committee. Fish welfare was prioritized throughout "
        "the experiment, with daily monitoring for signs of stress or disease."
    )
    
    doc.add_paragraph(
        "Humane handling practices were employed during all sampling procedures. Fish were "
        "anesthetized using clove oil (50 mg/L) before weighing and measuring to minimize "
        "stress. Any fish showing signs of severe illness or distress were humanely euthanized "
        "using an overdose of clove oil (200 mg/L)."
    )
    
    # ========== REFERENCES ==========
    add_heading(doc, 'References', 2)
    
    doc.add_paragraph(
        "AOAC International. (2023). Official methods of analysis of AOAC International "
        "(G. W. Latimer Jr., Ed.; 22nd ed.). Oxford University Press. "
        "https://doi.org/10.1093/9780197610145.001.0001"
    )
    
    doc.add_paragraph(
        "Kasihmuddin, S. M., Ghaffar, M. A., & Das, S. K. (2021). Rising temperature effects "
        "on growth and gastric emptying time of freshwater African catfish (Clarias gariepinus) "
        "fingerlings. Animals, 11(12), 3497. https://doi.org/10.3390/ani11123497"
    )
    
    doc.add_paragraph(
        "Obirikorang, K. A., Adjei-Boateng, D., Madkour, H. A., Otchere, F. A., & Skov, P. V. "
        "(2024). Nutritional requirements and effect of culture conditions on the performance "
        "of the African catfish (Clarias gariepinus): A review. Reviews in Aquaculture, 16(1), 1-25."
    )
    
    doc.add_paragraph(
        "Okomoda, V. T., Musa, S. O., Tiamiyu, L. O., Solomon, S. G., Ikhwanuddin, M., & "
        "Abol-Munafi, A. B. (2022). Biological performance of African catfish (Clarias gariepinus) "
        "fed varying feeding rates. Aquaculture Reports, 23, 101067."
    )
    
    # ========== DIAGRAM PLACEHOLDERS ==========
    add_heading(doc, 'Diagram Placeholders', 2)
    
    doc.add_paragraph(
        "The following diagrams should be exported from Mermaid Live Editor (https://mermaid.live) "
        "as PNG images and inserted at the indicated locations in the document:"
    )
    
    doc.add_paragraph("• Diagram 1: System Architecture")
    doc.add_paragraph("• Diagram 2: Feed Management Data Flow")
    doc.add_paragraph("• Diagram 3: Cost Effectiveness Analysis Flowchart")
    doc.add_paragraph("• Diagram 4: Labor and Time Comparison")
    doc.add_paragraph("• Diagram 5: Feed Wastage Monitoring System")
    doc.add_paragraph("• Diagram 6: Experimental Design for Feed Management Study")
    doc.add_paragraph("• Diagram 7: Hardware Block Diagram")
    doc.add_paragraph("• Diagram 8: Feed Dispensing Mechanism Flowchart")
    doc.add_paragraph("• Diagram 9: Economic Analysis Decision Tree")
    doc.add_paragraph("• Diagram 10: Software Architecture")
    doc.add_paragraph("• Diagram 11: Data Collection Timeline")
    doc.add_paragraph("• Diagram 12: Hardware Wiring Diagram")
    doc.add_paragraph("• Diagram 13: Cost-Benefit Analysis Process")
    
    # Save document
    output_path = 'docs/CHAPTER_3_FRIEND_FEED_MANAGEMENT_EXPANDED.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print("\nDocument generation complete!")
    print("\nNext steps:")
    print("1. Open the document in Microsoft Word")
    print("2. Export diagrams from Mermaid Live Editor as PNG")
    print("3. Insert diagrams at the indicated placeholder locations")
    print("4. Review and adjust formatting as needed")
    print("5. Fill in placeholders like [LOCATION], [DATE], etc.")

if __name__ == "__main__":
    main()

"""
Script to generate corrected Chapter 3 Word document with:
1. Nutrient Utilization Parameters section added
2. APA-style tables (3 horizontal lines only)
3. In-text references to all tables and figures
4. More specific section headings
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_apa_table_borders(table):
    """Apply APA-style borders: top, header-bottom, and table-bottom only"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Remove all borders first
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        if border_name in ['top', 'bottom']:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # 1.5pt
            border.set(qn('w:color'), '000000')
        else:
            border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # Add bottom border to header row
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:color'), '000000')
            tcBorders.append(bottom)
            tcPr.append(tcBorders)

def add_heading(doc, text, level):
    """Add a heading with proper formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para

def add_formula(doc, formula_text, description=""):
    """Add a formula with proper formatting"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(formula_text)
    run.italic = True
    if description:
        doc.add_paragraph(description)

def create_apa_table(doc, headers, data, caption=""):
    """Create an APA-style table with 3 horizontal lines only"""
    if caption:
        cap_para = doc.add_paragraph()
        cap_run = cap_para.add_run(caption)
        cap_run.bold = True
    
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    
    # Add data
    for row_idx, row_data in enumerate(data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
    
    # Apply APA borders
    set_apa_table_borders(table)
    
    doc.add_paragraph()  # Space after table
    return table

def main():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Title
    title = doc.add_heading('CHAPTER 3', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('MATERIALS AND METHODS', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =========================================================================
    # 3.1 Study Location and Experimental Tank Configuration
    # =========================================================================
    add_heading(doc, '3.1 Study Location and Experimental Tank Configuration', 2)
    
    doc.add_paragraph(
        'This study was conducted at [INSERT LAB NAME], Department of [INSERT DEPARTMENT], '
        '[INSERT INSTITUTION NAME], [INSERT LOCATION/CITY], Nigeria. The experiment was carried out '
        'to evaluate the effect of a smart assisted automatic fish feeder on the growth performance '
        'and nutrient utilization of African catfish (Clarias gariepinus). The experiment was designed '
        'as a comparative study between conventional manual feeding and smart automated feeding using '
        'a Q10 temperature-adjusted algorithm.'
    )
    
    doc.add_paragraph(
        'The experimental setup, as illustrated in Figure 3.1, consisted of two treatment groups: '
        'a control group that received manual feeding at fixed intervals and a treatment group that '
        'received automated feeding adjusted based on real-time water temperature measurements using '
        'the Q10 metabolic model.'
    )
    
    # Figure placeholder
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.1: Experimental Setup Layout]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.1: Experimental Setup Layout')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =========================================================================
    # 3.2 Materials and Equipment
    # =========================================================================
    add_heading(doc, '3.2 Materials and Equipment', 2)
    
    # 3.2.1 Hardware Components
    add_heading(doc, '3.2.1 Smart Feeder Hardware Components and Specifications', 3)
    
    doc.add_paragraph(
        'The smart assisted automatic fish feeder system was designed and constructed using the '
        'components listed in Table 3.1.'
    )
    
    # Table 3.1: Hardware Components
    hardware_headers = ['Component', 'Specification', 'Quantity', 'Purpose']
    hardware_data = [
        ['LILYGO T-A7670 R2', 'ESP32 + A7670G LTE', '1', 'Main controller with cellular connectivity'],
        ['ESP32-CAM-MB', 'OV2640 2MP Camera', '1', 'Feeding image capture and verification'],
        ['NEMA 23 Stepper Motor', '1.2 N·m, 2.8A, 1.8°, 200 steps/rev', '1', 'Auger drive motor'],
        ['DM542 Stepper Driver', '4.2A, 20-50V DC', '1', 'Motor control with microstepping'],
        ['20mm Wood Drill Auger', 'Stainless steel, 20mm pitch', '1', 'Feed dispensing (25g/rev)'],
        ['HX711 + 20kg Load Cell', '24-bit ADC', '1', 'Feed weight measurement'],
        ['DS18B20', 'Waterproof, 1M cable', '1', 'Water temperature sensing'],
        ['JSN-SR04T', 'Waterproof ultrasonic', '1', 'Water level detection'],
        ['Solar Panel', '50-100W monocrystalline', '1', 'Primary power source'],
        ['Charge Controller', 'MPPT/PWM 10-20A', '1', 'Battery charging'],
        ['12V Battery', '12-20Ah Lead-acid/LiFePO4', '1', 'Backup power'],
        ['LM2596 DC-DC', 'Adjustable step-down', '2', 'Voltage regulation (5V and 3.3V)'],
    ]
    create_apa_table(doc, hardware_headers, hardware_data, 'Table 3.1: Hardware Components Specification')

    # 3.2.2 Component Descriptions
    add_heading(doc, '3.2.2 Detailed Component Descriptions and Functions', 3)
    
    components = [
        ('LILYGO T-A7670 R2 Microcontroller', 
         'The LILYGO T-A7670 R2 served as the main controller for the smart fish feeder system. '
         'This development board integrated an ESP32-WROVER-B microcontroller with an A7670G 4G LTE Cat1 modem, '
         'enabling cellular connectivity in remote pond locations without WiFi access. The board featured a '
         'built-in 18650 battery holder with charging circuit, making it suitable for solar-powered applications. '
         'The ESP32 provided dual-core processing at 240MHz, 520KB SRAM, and multiple GPIO pins for interfacing '
         'with sensors and actuators.'),
        
        ('ESP32-CAM Module',
         'The ESP32-CAM (AI-Thinker) module was employed for feeding verification through image capture. '
         'This module featured an OV2640 camera sensor capable of capturing 2-megapixel JPEG images at '
         'resolutions up to 1600×1200 pixels. The camera was positioned to capture images of the feeding area '
         'during each dispensing event, providing visual verification of feed delivery.'),
        
        ('NEMA 23 Stepper Motor',
         'A NEMA 23 bipolar stepper motor with 1.2 N·m (1200 mN·m, approximately 170 oz-in) holding torque was selected '
         'as the prime mover for the feed dispensing mechanism. The motor featured a step angle of 1.8 degrees (200 steps '
         'per revolution), rated current of 2.8A per phase, and sufficient holding torque to drive the 20mm auger under load. '
         'The motor was operated at 30 RPM (800 steps per second with 8× microstepping), providing smooth rotation and '
         'precise feed quantity control.'),
        
        ('DM542 Stepper Motor Driver',
         'The DM542 digital stepper driver was used to control the NEMA 23 motor. This driver accepted step and '
         'direction pulse signals from the ESP32 microcontroller and generated the appropriate current waveforms '
         'for the motor coils. Microstepping was set to 8× (1600 steps per revolution) to achieve smoother motor '
         'operation and reduced vibration.'),
        
        ('20mm Wood Drill Auger',
         'A 20mm diameter stainless steel wood drill auger bit was adapted as the feed dispensing mechanism. '
         'Through calibration, it was determined that one revolution of the auger dispensed approximately 25 grams '
         'of commercial catfish pellets. This volumetric dispensing approach provided consistent feed delivery '
         'regardless of hopper fill level.'),
        
        ('HX711 Load Cell Amplifier and 20kg Load Cell',
         'The HX711 24-bit analog-to-digital converter (ADC) was paired with a 20kg capacity strain gauge load cell '
         'for precise weight measurement of the feed hopper. The HX711 provided high-resolution weight readings with '
         'a sensitivity of approximately 0.1 grams after calibration. The system generated low-feed alerts when the '
         'hopper weight dropped below 10% of capacity.'),
        
        ('DS18B20 Waterproof Temperature Sensor',
         'The DS18B20 digital temperature sensor was used to monitor water temperature in real-time. This sensor '
         'featured a waterproof stainless steel probe with 1-meter cable length. The sensor provided 12-bit resolution '
         '(0.0625°C precision) with an accuracy of ±0.5°C. Temperature readings were taken every 5 minutes and used '
         'by the Q10 algorithm to adjust feeding rates based on fish metabolic activity.'),
        
        ('JSN-SR04T Waterproof Ultrasonic Sensor',
         'The JSN-SR04T waterproof ultrasonic distance sensor was employed for water level monitoring. This sensor '
         'featured an IP67-rated transducer probe suitable for outdoor and wet environments, with an effective range '
         'of 25cm to 400cm and accuracy of ±1cm.'),
        
        ('Solar Panel and Power System',
         'A monocrystalline solar panel rated at 50-80W served as the primary power source for the smart feeder system. '
         'A 12V rechargeable battery with 12-20Ah capacity provided energy storage for operation during nighttime and '
         'cloudy periods. The battery capacity was sized to provide at least 24 hours of autonomous operation without '
         'solar charging.'),
    ]
    
    for title, description in components:
        para = doc.add_paragraph()
        run = para.add_run(f'{title}: ')
        run.bold = True
        para.add_run(description)
    
    # Figure 3.2 placeholder
    doc.add_paragraph()
    doc.add_paragraph('The hardware block diagram showing the interconnection of all components is presented in Figure 3.2.')
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.2: Hardware Block Diagram]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.2: Hardware Block Diagram')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3.2.3 Sensor Specifications
    add_heading(doc, '3.2.3 Environmental Sensor Specifications', 3)
    
    doc.add_paragraph('The specifications of all sensors used in the system are detailed in Table 3.2.')
    
    sensor_headers = ['Sensor', 'Model', 'Range', 'Accuracy', 'Interface', 'Purpose']
    sensor_data = [
        ['Temperature', 'DS18B20', '-55 to 125°C', '±0.5°C', 'OneWire', 'Water temperature'],
        ['Weight', 'HX711 + Load Cell', '0-20kg', '±0.1g', 'Digital SPI', 'Feed level'],
        ['Water Level', 'JSN-SR04T', '25-400cm', '±1cm', 'Ultrasonic', 'Water level'],
        ['Camera', 'OV2640', '2MP JPEG', 'VGA/SVGA', 'UART', 'Image capture'],
    ]
    create_apa_table(doc, sensor_headers, sensor_data, 'Table 3.2: Sensor Specifications')
    
    # 3.2.4 Motor Control System
    add_heading(doc, '3.2.4 Stepper Motor Control System Configuration', 3)
    
    doc.add_paragraph(
        'The feed dispensing mechanism utilized a NEMA 23 bipolar stepper motor coupled with a DM542 stepper driver '
        'and a 20mm wood drill auger. The motor control parameters configured for precise feed dispensing are shown '
        'in Table 3.3.'
    )
    
    motor_headers = ['Parameter', 'Value', 'Unit', 'Description']
    motor_data = [
        ['Motor Type', 'NEMA 23', '-', 'Bipolar stepper motor'],
        ['Holding Torque', '1.2', 'N·m', 'Equivalent to 1200 mN·m or ~170 oz-in'],
        ['Steps per Revolution', '200', 'steps', '1.8 degrees per step'],
        ['Microstepping', '8', 'microsteps', 'Set via DM542 DIP switches'],
        ['Effective Steps per Rev', '1600', 'steps', '200 × 8 microsteps'],
        ['Motor Current', '2.8', 'Amperes', 'Peak current per phase'],
        ['Pulse Width', '5', 'microseconds', 'Minimum HIGH time for step pulse'],
        ['Step Delay', '1250', 'microseconds', 'Time between steps at 800 steps/sec'],
        ['Max Speed', '800', 'steps/second', 'Optimized for 1.2 N·m torque'],
        ['Motor RPM', '30', 'RPM', '(800 × 60) / 1600 steps'],
        ['Grams per Revolution', '25', 'grams', 'Calibrated feed output'],
        ['Dispensing Rate', '750', 'grams/minute', '30 RPM × 25g per rev'],
    ]
    create_apa_table(doc, motor_headers, motor_data, 'Table 3.3: Motor Control Parameters')
    
    doc.add_paragraph('The timing diagram for motor control signals is illustrated in Figure 3.3.')
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.3: Motor Control Signal Timing Diagram]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.3: Motor Control Signal Timing Diagram')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # DIP Switch Configuration
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('DM542 DIP Switch Configuration:')
    run.bold = True
    doc.add_paragraph('For 8 microstepping (1600 steps/rev): SW1=OFF, SW2=ON, SW3=ON')
    doc.add_paragraph('For 2.8A motor current: SW4=ON, SW5=OFF, SW6=ON')
    
    # 3.2.5 Software Technologies
    add_heading(doc, '3.2.5 Software Development Technologies and Frameworks', 3)
    
    doc.add_paragraph(
        'The smart fish feeder system comprised three software layers: embedded firmware running on the ESP32 '
        'microcontroller, a cloud backend providing API services and data storage, and a mobile application for '
        'user interaction and monitoring. The software technologies used across all three layers of the system '
        'are summarized in Table 3.4.'
    )
    
    software_headers = ['Layer', 'Technology', 'Version', 'Purpose']
    software_data = [
        ['Firmware', 'PlatformIO', '6.x', 'ESP32 development environment'],
        ['Firmware', 'Arduino Framework', '2.x', 'Hardware abstraction layer'],
        ['Firmware', 'TinyGSM', '0.11.x', 'LTE modem communication'],
        ['Backend', 'Go', '1.21+', 'Server-side programming'],
        ['Backend', 'Gin', '1.9+', 'HTTP web framework'],
        ['Backend', 'PostgreSQL', '15+', 'Relational database'],
        ['Backend', 'Redis', '7+', 'Caching and sessions'],
        ['Mobile', 'Flutter', '3.16+', 'Cross-platform UI framework'],
        ['Mobile', 'Dart', '3.2+', 'Programming language'],
        ['Mobile', 'Riverpod', '2.4+', 'State management'],
    ]
    create_apa_table(doc, software_headers, software_data, 'Table 3.4: Software Technologies')
    
    doc.add_paragraph(
        'The software architecture diagram depicting the relationship between firmware, backend, and mobile layers '
        'is shown in Figure 3.4.'
    )
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.4: Software Architecture Diagram]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.4: Software Architecture Diagram')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================================
    # 3.3 Experimental Methods and Procedures
    # =========================================================================
    add_heading(doc, '3.3 Experimental Methods and Procedures', 2)
    
    # 3.3.1 System Architecture
    add_heading(doc, '3.3.1 Three-Layer System Architecture Design', 3)
    
    doc.add_paragraph(
        'The smart fish feeder system architecture consisted of three main layers: the hardware layer '
        '(sensors and actuators at the fish pond), the cloud infrastructure (backend services, database, '
        'and MQTT broker), and the mobile application layer (user interface for monitoring and control). '
        'Communication between the hardware layer and cloud infrastructure was established via LTE cellular '
        'network using MQTT protocol, while the mobile application communicated with the backend via HTTPS REST API. '
        'The complete system architecture, including hardware, cloud, and mobile layers, is presented in Figure 3.5.'
    )
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.5: System Architecture Diagram]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.5: System Architecture Diagram')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3.3.2 Q10 Algorithm
    add_heading(doc, '3.3.2 Q10 Temperature-Adjusted Metabolic Feeding Algorithm', 3)
    
    doc.add_paragraph(
        'The Q10 model was employed to adjust feeding rates based on water temperature. The Q10 coefficient '
        'represented the factor by which metabolic rate increased for every 10°C rise in temperature. For fish, '
        'Q10 values typically range from 2.0 to 2.5, with an average of approximately 2.0 (Kasihmuddin et al., 2021; '
        'Obirikorang et al., 2024). This temperature-dependent feeding approach aimed to match feed supply with the metabolic '
        'demands of the fish, potentially improving feed conversion efficiency and reducing feed waste. '
        'The Q10 algorithm parameters used for temperature-adjusted feeding are listed in Table 3.5.'
    )
    
    # Q10 Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Q10 Metabolic Adjustment Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('FR_adj = SFR × Q10^((T - T_ref) / 10) × TIF').italic = True
    
    doc.add_paragraph('Where:')
    doc.add_paragraph('• FR_adj = Adjusted feeding rate (g/day)')
    doc.add_paragraph('• SFR = Standard feeding rate (g/day)')
    doc.add_paragraph('• Q10 = Temperature coefficient (2.0 for fish, typical range 2.0-2.5)')
    doc.add_paragraph('• T = Current water temperature (°C)')
    doc.add_paragraph('• T_ref = Reference temperature (25°C)')
    doc.add_paragraph('• TIF = Thermal Inhibition Factor (penalty when T > 30°C or T < 20°C)')
    
    q10_headers = ['Parameter', 'Symbol', 'Value', 'Unit', 'Description']
    q10_data = [
        ['Q10 Coefficient', 'Q10', '2.0', 'dimensionless', 'Metabolic rate temperature sensitivity'],
        ['Reference Temperature', 'Tref', '25', '°C', 'Standard temperature for calculations'],
        ['Optimal Temperature Range', 'Topt', '25-30', '°C', 'Best growth temperature'],
        ['Critical Low Temperature', 'Tcrit_low', '20', '°C', 'Feeding reduced below this'],
        ['Critical High Temperature', 'Tcrit_high', '30', '°C', 'Feeding reduced above this'],
        ['Thermal Inhibition Factor', 'TIF', '0.3-0.8', 'dimensionless', 'Penalty during thermal stress'],
        ['Standard Feeding Rate', 'SFR', '3-5', '% BW/day', 'Base feeding rate by body weight'],
    ]
    create_apa_table(doc, q10_headers, q10_data, 'Table 3.5: Q10 Algorithm Parameters')
    
    doc.add_paragraph(
        'The feeding control algorithm was divided into two phases: the decision phase (Figure 3.6a) and '
        'the execution phase (Figure 3.6b).'
    )
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.6a: Feeding Control Algorithm - Decision Phase]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.6a: Feeding Control Algorithm - Decision Phase')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.6b: Feeding Control Algorithm - Execution Phase]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.6b: Feeding Control Algorithm - Execution Phase')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3.3.3 Feed Dispensing Process
    add_heading(doc, '3.3.3 Automated Feed Dispensing Process and Calculations', 3)
    
    doc.add_paragraph(
        'The feed dispensing process involved calculating the required motor steps based on the target feed amount, '
        'executing the motor movement, and verifying the dispensed quantity using the load cell sensor. The process '
        'was initiated either by scheduled feeding times or manual commands from the mobile application.'
    )
    
    # Formulas
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Motor Steps Calculation:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('Total_Steps = (Target_Grams / Grams_Per_Rev) × Steps_Per_Rev × Microsteps').italic = True
    
    para = doc.add_paragraph()
    run = para.add_run('Motor RPM Calculation:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('RPM = (Steps_Per_Second × 60) / Effective_Steps_Per_Revolution').italic = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('RPM = (800 × 60) / 1600 = 30 RPM').italic = True
    
    para = doc.add_paragraph()
    run = para.add_run('Feed Dispensing Rate:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('Dispensing_Rate = RPM × Grams_Per_Revolution = 30 × 25 = 750 grams/minute').italic = True
    
    doc.add_paragraph('An example calculation for feed dispensing is provided in Table 3.6.')
    
    calc_headers = ['Step', 'Calculation', 'Value', 'Unit']
    calc_data = [
        ['Target Feed Amount', 'User input', '100', 'grams'],
        ['Q10 Adjustment Factor', 'Q10^((T-25)/10)', '1.15', 'dimensionless'],
        ['Adjusted Feed Amount', '100 × 1.15', '115', 'grams'],
        ['Revolutions Needed', '115 / 25', '4.6', 'revolutions'],
        ['Steps per Revolution', '200 × 8', '1600', 'steps'],
        ['Total Steps Required', '4.6 × 1600', '7360', 'steps'],
        ['Time per Step', '1 / 800', '1.25', 'milliseconds'],
        ['Total Dispensing Time', '7360 × 1.25', '9.2', 'seconds'],
    ]
    create_apa_table(doc, calc_headers, calc_data, 'Table 3.6: Feed Dispensing Calculation Example')
    
    doc.add_paragraph(
        'The feed dispensing mechanism flowchart is presented in two parts: the calculation phase (Figure 3.7a) '
        'and the motor execution phase (Figure 3.7b).'
    )
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.7a: Feed Dispensing Mechanism - Calculation Phase]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.7a: Feed Dispensing Mechanism - Calculation Phase')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.7b: Feed Dispensing Mechanism - Motor Execution Phase]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.7b: Feed Dispensing Mechanism - Motor Execution Phase')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 3.3.4 Experimental Design
    add_heading(doc, '3.3.4 Completely Randomized Experimental Design', 3)
    
    doc.add_paragraph(
        'A completely randomized design (CRD) was employed with two treatment groups. African catfish fingerlings '
        'were randomly assigned to either the control group (manual feeding) or the treatment group (smart automated '
        'feeding). The experiment was conducted over an 8-week period with weekly measurements of growth parameters. '
        'The experimental design parameters for both control and treatment groups are summarized in Table 3.7.'
    )
    
    exp_headers = ['Parameter', 'Control Group', 'Treatment Group']
    exp_data = [
        ['Feeding Method', 'Manual', 'Smart Automatic Feeder'],
        ['Feeding Schedule', 'Fixed 3× daily', 'Q10-adjusted dynamic'],
        ['Feed Amount', '3% body weight fixed', 'Temperature-adjusted'],
        ['Temperature Monitoring', 'Manual thermometer', 'Continuous DS18B20 sensor'],
        ['Data Recording', 'Manual logbook', 'Automatic cloud logging'],
        ['Feeding Verification', 'Visual observation', 'ESP32-CAM image capture'],
        ['Number of Fish', '50', '50'],
        ['Tank Size', '500 liters', '500 liters'],
        ['Duration', '8 weeks', '8 weeks'],
        ['Feed Type', 'Commercial catfish pellets', 'Commercial catfish pellets'],
        ['Water Exchange', '30% weekly', '30% weekly'],
    ]
    create_apa_table(doc, exp_headers, exp_data, 'Table 3.7: Experimental Design Parameters')
    
    doc.add_paragraph(
        'The experimental design flowchart showing the progression from study design to statistical analysis '
        'is illustrated in Figure 3.8.'
    )
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.8: Experimental Design Flowchart]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.8: Experimental Design Flowchart')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.3.5 Growth Performance Parameters
    add_heading(doc, '3.3.5 Growth Performance Parameter Measurements', 3)
    
    doc.add_paragraph(
        'The following growth performance parameters were measured and calculated to evaluate the effect of the '
        'smart feeding system on African catfish growth. These parameters are defined in Table 3.8.'
    )
    
    growth_headers = ['Parameter', 'Formula', 'Unit', 'Description']
    growth_data = [
        ['Initial Weight (W₁)', 'Direct measurement', 'g', 'Weight at start of experiment'],
        ['Final Weight (W₂)', 'Direct measurement', 'g', 'Weight at end of experiment'],
        ['Weight Gain (WG)', 'WG = W₂ - W₁', 'g', 'Total weight increase'],
        ['Average Daily Gain (ADG)', 'ADG = WG / Days', 'g/day', 'Daily weight increase'],
        ['Specific Growth Rate (SGR)', 'SGR = ((ln W₂ - ln W₁) / t) × 100', '%/day', 'Daily percentage growth'],
        ['Condition Factor (K)', 'K = (W / L³) × 100', '-', 'Body condition index'],
        ['Survival Rate (SR)', 'SR = (Final Count / Initial Count) × 100', '%', 'Percentage surviving'],
    ]
    create_apa_table(doc, growth_headers, growth_data, 'Table 3.8: Growth Performance Parameters')
    
    # SGR Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Specific Growth Rate (SGR) Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('SGR = ((ln W₂ - ln W₁) / t) × 100').italic = True
    
    doc.add_paragraph('Where:')
    doc.add_paragraph('• W₂ = Final body weight (g)')
    doc.add_paragraph('• W₁ = Initial body weight (g)')
    doc.add_paragraph('• t = Time period (days)')
    doc.add_paragraph('• ln = Natural logarithm')
    
    # =========================================================================
    # 3.3.6 NUTRIENT UTILIZATION PARAMETERS - NEW SECTION
    # =========================================================================
    add_heading(doc, '3.3.6 Nutrient Utilization Parameter Measurements', 3)
    
    doc.add_paragraph(
        'The nutrient utilization parameters were measured and calculated to evaluate the efficiency of nutrient '
        'conversion in African catfish fed using the smart feeding system compared to manual feeding. These parameters, '
        'presented in Table 3.9, were critical for assessing the economic efficiency of the feeding system and its '
        'impact on nutrient utilization as stated in the study objectives.'
    )
    
    nutrient_headers = ['Parameter', 'Formula', 'Unit', 'Description']
    nutrient_data = [
        ['Feed Conversion Ratio (FCR)', 'FCR = Feed Intake / Weight Gain', 'ratio', 'Feed efficiency measure'],
        ['Protein Efficiency Ratio (PER)', 'PER = Weight Gain / Protein Intake', 'ratio', 'Protein utilization efficiency'],
        ['Apparent Net Protein Utilization (ANPU)', 'ANPU = ((Pf - Pi) / Protein Intake) × 100', '%', 'Protein retention percentage'],
        ['Protein Productive Value (PPV)', 'PPV = (Protein Gain / Protein Intake) × 100', '%', 'Protein conversion efficiency'],
        ['Lipid Efficiency Ratio (LER)', 'LER = Weight Gain / Lipid Intake', 'ratio', 'Lipid utilization efficiency'],
        ['Lipid Retention Efficiency (LRE)', 'LRE = ((Lf - Li) / Lipid Intake) × 100', '%', 'Lipid retention percentage'],
        ['Gross Energy Retention (GER)', 'GER = ((Ef - Ei) / Energy Intake) × 100', '%', 'Energy retention percentage'],
    ]
    create_apa_table(doc, nutrient_headers, nutrient_data, 'Table 3.9: Nutrient Utilization Parameters')
    
    # FCR Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Feed Conversion Ratio (FCR) Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('FCR = Total Feed Consumed (g) / Total Weight Gain (g)').italic = True
    
    doc.add_paragraph('Note: Lower FCR indicated better feed efficiency. Target FCR for African catfish under optimal conditions: 1.5 - 2.0 (Okomoda et al., 2022).')
    
    # PER Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Protein Efficiency Ratio (PER) Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('PER = Weight Gain (g) / Protein Intake (g)').italic = True
    
    doc.add_paragraph('Where:')
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('Protein Intake = Feed Consumed × (% Crude Protein in Feed / 100)').italic = True
    
    # ANPU Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Apparent Net Protein Utilization (ANPU) Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('ANPU (%) = ((Final Body Protein - Initial Body Protein) / Protein Intake) × 100').italic = True
    
    doc.add_paragraph('Where:')
    doc.add_paragraph('• Pf = Final body protein content (g)')
    doc.add_paragraph('• Pi = Initial body protein content (g)')
    
    # LER Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Lipid Efficiency Ratio (LER) Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('LER = Weight Gain (g) / Lipid Intake (g)').italic = True
    
    doc.add_paragraph('Where:')
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('Lipid Intake = Feed Consumed × (% Crude Lipid in Feed / 100)').italic = True
    
    # LRE Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Lipid Retention Efficiency (LRE) Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('LRE (%) = ((Final Body Lipid - Initial Body Lipid) / Lipid Intake) × 100').italic = True
    
    # GER Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Gross Energy Retention (GER) Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('GER (%) = ((Final Body Energy - Initial Body Energy) / Energy Intake) × 100').italic = True
    
    doc.add_paragraph('Where:')
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('Energy Intake = Feed Consumed × Gross Energy Content of Feed (kJ/g)').italic = True
    
    # Proximate Composition Analysis
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('Proximate Composition Analysis:')
    run.bold = True
    
    doc.add_paragraph(
        'To calculate the nutrient utilization parameters, proximate composition analysis was performed on fish '
        'carcasses at the beginning and end of the experiment following standard methods (AOAC, 2023). A representative '
        'sample of fish from each treatment group was sacrificed and analyzed for crude protein, crude lipid, moisture '
        'content, ash content, and gross energy. The analysis methods are summarized in Table 3.10.'
    )
    
    prox_headers = ['Parameter', 'Method', 'Reference']
    prox_data = [
        ['Crude Protein', 'Kjeldahl (N × 6.25)', 'AOAC 2023'],
        ['Crude Lipid', 'Soxhlet Extraction', 'AOAC 2023'],
        ['Moisture', 'Oven Drying (105°C)', 'AOAC 2023'],
        ['Ash', 'Incineration (550°C)', 'AOAC 2023'],
        ['Gross Energy', 'Bomb Calorimetry', 'AOAC 2023'],
    ]
    create_apa_table(doc, prox_headers, prox_data, 'Table 3.10: Proximate Composition Analysis Methods')

    # 3.3.7 Statistical Analysis
    add_heading(doc, '3.3.7 Statistical Analysis Using Independent Samples t-Test', 3)
    
    doc.add_paragraph(
        'Data collected from the experiment were analyzed using descriptive and inferential statistics. '
        'The Independent Samples t-Test was employed to compare the means of growth performance and nutrient '
        'utilization parameters between the control group (manual feeding) and the treatment group (smart automated '
        'feeding). The statistical analysis methods employed in this study are outlined in Table 3.11.'
    )
    
    para = doc.add_paragraph()
    run = para.add_run('Independent Samples t-Test:')
    run.bold = True
    
    doc.add_paragraph(
        'The t-Test was chosen as the appropriate statistical method because the study involved comparing the means '
        'of two independent groups (control vs. treatment). The following assumptions were verified before conducting '
        'the t-Test:'
    )
    
    doc.add_paragraph('1. Independence of observations: Fish were randomly assigned to groups')
    doc.add_paragraph('2. Normal distribution: Data were tested for normality using Shapiro-Wilk test')
    doc.add_paragraph('3. Homogeneity of variance: Levene\'s test was used to verify equal variances')
    
    # t-Test Formula
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run('t-Test Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('t = (X̄₁ - X̄₂) / √(s²p × (1/n₁ + 1/n₂))').italic = True
    
    doc.add_paragraph('Where:')
    doc.add_paragraph('• X̄₁, X̄₂ = Sample means of group 1 and group 2')
    doc.add_paragraph('• s²p = Pooled variance')
    doc.add_paragraph('• n₁, n₂ = Sample sizes of group 1 and group 2')
    
    para = doc.add_paragraph()
    run = para.add_run('Pooled Variance Formula:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('s²p = ((n₁-1)s₁² + (n₂-1)s₂²) / (n₁ + n₂ - 2)').italic = True
    
    para = doc.add_paragraph()
    run = para.add_run('Degrees of Freedom:')
    run.bold = True
    
    formula_para = doc.add_paragraph()
    formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_para.add_run('df = n₁ + n₂ - 2').italic = True
    
    para = doc.add_paragraph()
    run = para.add_run('Statistical Significance:')
    run.bold = True
    
    doc.add_paragraph(
        'A significance level of α = 0.05 was used for all statistical tests. Results were considered statistically '
        'significant when p < 0.05. The null hypothesis (H₀) stated that there was no significant difference between '
        'the means of the two groups, while the alternative hypothesis (H₁) stated that there was a significant difference.'
    )
    
    doc.add_paragraph('• H₀: μ₁ = μ₂ (No significant difference between groups)')
    doc.add_paragraph('• H₁: μ₁ ≠ μ₂ (Significant difference exists between groups)')
    
    para = doc.add_paragraph()
    run = para.add_run('Decision Rule:')
    run.bold = True
    
    doc.add_paragraph('• If |t_calculated| > t_critical or p-value < 0.05: Reject H₀')
    doc.add_paragraph('• If |t_calculated| ≤ t_critical or p-value ≥ 0.05: Fail to reject H₀')
    
    stat_headers = ['Analysis', 'Method', 'Software', 'Purpose']
    stat_data = [
        ['Descriptive Statistics', 'Mean, SD, SE', 'Microsoft Excel / SPSS', 'Summarize data'],
        ['Normality Test', 'Shapiro-Wilk', 'SPSS', 'Verify normal distribution'],
        ['Variance Test', 'Levene\'s Test', 'SPSS', 'Verify homogeneity of variance'],
        ['Inferential Statistics', 'Independent t-Test', 'SPSS', 'Compare group means'],
        ['Significance Level', 'α = 0.05', '-', 'Determine statistical significance'],
    ]
    create_apa_table(doc, stat_headers, stat_data, 'Table 3.11: Statistical Analysis Methods')
    
    # 3.3.8 Hardware Wiring
    add_heading(doc, '3.3.8 Hardware Wiring and Circuit Connections', 3)
    
    doc.add_paragraph(
        'The following wiring connections were established between the microcontroller and peripheral components. '
        'The motor driver wiring connections are detailed in Table 3.12, while the sensor wiring connections are '
        'provided in Table 3.13.'
    )
    
    para = doc.add_paragraph()
    run = para.add_run('Motor Driver Connections (DM542):')
    run.bold = True
    
    motor_wire_headers = ['ESP32 Pin', 'DM542 Terminal', 'Function']
    motor_wire_data = [
        ['GPIO32', 'PUL+', 'Step pulse signal'],
        ['GPIO33', 'DIR+', 'Direction signal'],
        ['GPIO0', 'ENA+', 'Enable signal'],
        ['GND', 'PUL-, DIR-, ENA-', 'Common ground'],
        ['12-48V DC', 'Power terminals', 'Motor power supply'],
    ]
    create_apa_table(doc, motor_wire_headers, motor_wire_data, 'Table 3.12: Motor Driver Wiring')
    
    para = doc.add_paragraph()
    run = para.add_run('Sensor Connections:')
    run.bold = True
    
    sensor_wire_headers = ['Sensor', 'ESP32 Pin', 'Function']
    sensor_wire_data = [
        ['HX711 DOUT', 'GPIO39 (VN)', 'Load cell data output'],
        ['HX711 SCK', 'GPIO5', 'Load cell clock'],
        ['DS18B20 Data', 'GPIO23', 'Temperature data (with 4.7kΩ pullup)'],
        ['JSN-SR04T TRIG', 'GPIO17', 'Ultrasonic trigger'],
        ['JSN-SR04T ECHO', 'GPIO34', 'Ultrasonic echo'],
    ]
    create_apa_table(doc, sensor_wire_headers, sensor_wire_data, 'Table 3.13: Sensor Wiring Connections')
    
    doc.add_paragraph('The complete wiring diagram showing all electrical connections is presented in Figure 3.9.')
    
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run('[INSERT FIGURE 3.9: Complete Wiring Diagram]').italic = True
    
    cap_para = doc.add_paragraph('Figure 3.9: Complete Wiring Diagram')
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =========================================================================
    # 3.4 Data Collection
    # =========================================================================
    add_heading(doc, '3.4 Data Collection Methods and Storage', 2)
    
    doc.add_paragraph(
        'Data collection was performed automatically by the smart feeder system and manually for certain parameters. '
        'The following data were collected throughout the experiment:'
    )
    
    doc.add_paragraph('• Water temperature: Recorded every 5 minutes via DS18B20 sensor')
    doc.add_paragraph('• Feed dispensed: Recorded for each feeding event via HX711 load cell')
    doc.add_paragraph('• Feeding images: Captured via ESP32-CAM for each feeding event')
    doc.add_paragraph('• Fish weight: Measured weekly using digital scale (individual sampling)')
    doc.add_paragraph('• Fish length: Measured weekly using measuring board')
    doc.add_paragraph('• Mortality: Recorded daily')
    doc.add_paragraph('• Water quality: Monitored weekly (pH, ammonia, nitrite)')
    doc.add_paragraph('• Proximate composition: Analyzed at beginning and end of experiment')
    
    doc.add_paragraph(
        'All sensor data were transmitted via LTE cellular connection to the cloud backend and stored in a PostgreSQL '
        'database for analysis. The mobile application provided real-time monitoring and historical data visualization. '
        'Feeding images were stored on Cloudinary CDN and linked to feeding event records in the database.'
    )
    
    # =========================================================================
    # 3.5 Ethical Considerations
    # =========================================================================
    add_heading(doc, '3.5 Ethical Considerations in Animal Research', 2)
    
    doc.add_paragraph(
        'This study was conducted in accordance with ethical guidelines for animal research. The fish were handled '
        'humanely, and proper care was taken to minimize stress during measurements. Water quality was maintained '
        'within optimal ranges for African catfish, and any fish showing signs of disease or distress were promptly '
        'treated or removed from the experiment. Feeding rates were designed to meet nutritional requirements without '
        'overfeeding, which could lead to water quality deterioration and health issues.'
    )
    
    # =========================================================================
    # REFERENCES (APA 7th Edition)
    # =========================================================================
    doc.add_page_break()
    
    ref_heading = doc.add_heading('REFERENCES', 1)
    ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # List of references in APA 7th edition format (within 5 years: 2021-2026)
    references = [
        'AOAC International. (2023). Official methods of analysis of AOAC International (G. W. Latimer Jr., Ed.; 22nd ed.). Oxford University Press. https://doi.org/10.1093/9780197610145.001.0001',
        
        'Kasihmuddin, S. M., Ghaffar, M. A., & Das, S. K. (2021). Rising temperature effects on growth and gastric emptying time of freshwater African catfish (Clarias gariepinus) fingerlings. Animals, 11(12), 3497. https://doi.org/10.3390/ani11123497',
        
        'Obirikorang, K. A., Adjei-Boateng, D., Madkour, H. A., Otchere, F. A., & Skov, P. V. (2024). Nutritional requirements and effect of culture conditions on the performance of the African catfish (Clarias gariepinus): A review. Reviews in Aquaculture, 16(1), 1-25.',
        
        'Okomoda, V. T., Musa, S. O., Tiamiyu, L. O., Solomon, S. G., Ikhwanuddin, M., & Abol-Munafi, A. B. (2022). Biological performance of African catfish (Clarias gariepinus) fed varying feeding rates. Aquaculture Reports, 23, 101067.',
    ]
    
    for ref in references:
        para = doc.add_paragraph(ref)
        para.paragraph_format.first_line_indent = Cm(-1.27)  # Hanging indent
        para.paragraph_format.left_indent = Cm(1.27)
        para.paragraph_format.space_after = Pt(12)
    
    # Save document
    output_path = 'docs/CHAPTER_3_CORRECTED.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    print('\\nDocument includes:')
    print('1. ✓ Nutrient Utilization Parameters section (NEW - addresses supervisor feedback)')
    print('2. ✓ APA-style tables with 3 horizontal lines only')
    print('3. ✓ In-text references to all tables and figures')
    print('4. ✓ More specific section headings')
    print('5. ✓ Proximate composition analysis methods')
    print('6. ✓ APA 7th Edition References section')
    print('\\nPlaceholders for figures are marked with [INSERT FIGURE X.X: ...]')

if __name__ == '__main__':
    main()

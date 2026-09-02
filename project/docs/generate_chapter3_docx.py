"""
Generate Chapter 3: Materials and Methods Word Document (PROPOSAL - Past Tense)
With detailed component descriptions
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_image_placeholder(doc, caption, fig_num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n[INSERT FIGURE {fig_num} HERE]\n\n")
    run.font.size = Pt(12)
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(f"Figure {fig_num}: {caption}")
    run.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph()

def add_table(doc, headers, data, caption=None, table_num=None):
    if caption and table_num:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap_p.add_run(f"Table {table_num}: {caption}")
        run.bold = True
        run.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'D9E2F3')
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
    doc.add_paragraph()
    return table

def add_formula(doc, formula_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)

def create_chapter3_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # TITLE
    title = doc.add_heading('CHAPTER 3', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('MATERIALS AND METHODS', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.1 STUDY AREA
    add_heading(doc, '3.1 Study Area and Experimental Setup', level=2)
    doc.add_paragraph(
        'This study was conducted to evaluate the effect of a smart assisted automatic fish feeder '
        'on the growth performance and nutrient utilization of African catfish (Clarias gariepinus). '
        'The experiment was designed as a comparative study between conventional manual feeding '
        'and smart automated feeding using a Q10 temperature-adjusted algorithm.'
    )
    doc.add_paragraph(
        'The experimental setup consisted of two treatment groups: a control group that received '
        'manual feeding at fixed intervals and a treatment group that received automated feeding '
        'adjusted based on real-time water temperature measurements using the Q10 metabolic model.'
    )
    add_image_placeholder(doc, 'Experimental Setup Layout', '3.1')

    # 3.2 MATERIALS
    add_heading(doc, '3.2 Materials', level=2)
    add_heading(doc, '3.2.1 Hardware Components', level=3)
    doc.add_paragraph(
        'The smart fish feeder system was designed and constructed using the following components:'
    )

    hardware_headers = ['Component', 'Specification', 'Quantity', 'Purpose']
    hardware_data = [
        ['LILYGO T-A7670 R2', 'ESP32 + A7670G LTE', '1', 'Main controller with cellular connectivity'],
        ['ESP32-CAM-MB', 'OV2640 2MP Camera', '1', 'Feeding image capture and verification'],
        ['NEMA 23 Stepper Motor', '2.8A, 1.8°/step, 200 steps/rev', '1', 'Auger drive motor'],
        ['DM542 Stepper Driver', '4.2A, 20-50V DC', '1', 'Motor control with microstepping'],
        ['20mm Wood Drill Auger', 'Stainless steel, 20mm pitch', '1', 'Feed dispensing (25g/rev)'],
        ['HX711 + 20kg Load Cell', '24-bit ADC', '1', 'Feed weight measurement'],
        ['DS18B20', 'Waterproof, 1M cable', '1', 'Water temperature sensing'],
        ['JSN-SR04T', 'Waterproof ultrasonic, 25-400cm', '1', 'Water level detection'],
        ['Solar Panel', '50-100W monocrystalline', '1', 'Primary power source'],
        ['Charge Controller', 'MPPT/PWM 10-20A', '1', 'Battery charging'],
        ['12V Battery', '12-20Ah Lead-acid/LiFePO4', '1', 'Backup power storage'],
        ['LM2596 DC-DC', 'Adjustable step-down', '2', 'Voltage regulation (5V and 3.3V)'],
    ]
    add_table(doc, hardware_headers, hardware_data, 'Hardware Components Specification', '3.1')

    # =========================================================================
    # 3.2.2 DETAILED COMPONENT DESCRIPTIONS
    # =========================================================================
    add_heading(doc, '3.2.2 Detailed Component Descriptions', level=3)

    # LILYGO T-A7670 R2
    p = doc.add_paragraph()
    run = p.add_run('LILYGO T-A7670 R2 Microcontroller: ')
    run.bold = True
    p.add_run(
        'The LILYGO T-A7670 R2 served as the main controller for the smart fish feeder system. '
        'This development board integrated an ESP32-WROVER-B microcontroller with an A7670G 4G LTE '
        'Cat1 modem, enabling cellular connectivity in remote pond locations without WiFi access. '
        'The board featured a built-in 18650 battery holder with charging circuit, making it suitable '
        'for solar-powered applications. The ESP32 provided dual-core processing at 240MHz, 520KB SRAM, '
        'and multiple GPIO pins for interfacing with sensors and actuators. The A7670G modem supported '
        'LTE-FDD/LTE-TDD/GSM/GPRS/EDGE networks and enabled MQTT communication with the cloud backend '
        'for real-time data transmission and remote control.'
    )

    # ESP32-CAM
    p = doc.add_paragraph()
    run = p.add_run('ESP32-CAM Module: ')
    run.bold = True
    p.add_run(
        'The ESP32-CAM (AI-Thinker) module was employed for feeding verification through image capture. '
        'This module featured an OV2640 camera sensor capable of capturing 2-megapixel JPEG images at '
        'resolutions up to 1600×1200 pixels. The camera was positioned to capture images of the feeding '
        'area during each dispensing event, providing visual verification of feed delivery and enabling '
        'future analysis of fish feeding behavior. The module communicated with the main controller via '
        'UART serial interface at 115200 baud rate. Images were compressed as JPEG and uploaded to '
        'Cloudinary CDN through the cloud backend for storage and retrieval via the mobile application.'
    )

    # NEMA 23 Motor
    p = doc.add_paragraph()
    run = p.add_run('NEMA 23 Stepper Motor: ')
    run.bold = True
    p.add_run(
        'A NEMA 23 bipolar stepper motor was selected as the prime mover for the feed dispensing mechanism. '
        'The motor featured a step angle of 1.8 degrees (200 steps per revolution), rated current of 2.8A '
        'per phase, and holding torque sufficient to drive the auger under load. The NEMA 23 frame size '
        '(57mm × 57mm) provided adequate torque for dispensing fish feed pellets through the auger mechanism. '
        'The motor was operated at 30 RPM (800 steps per second with 8× microstepping), providing smooth '
        'rotation and precise feed quantity control. The stepper motor design ensured accurate positioning '
        'without the need for feedback sensors, as each step pulse corresponded to a known angular displacement.'
    )

    # DM542 Driver
    p = doc.add_paragraph()
    run = p.add_run('DM542 Stepper Motor Driver: ')
    run.bold = True
    p.add_run(
        'The DM542 digital stepper driver was used to control the NEMA 23 motor. This driver accepted '
        'step and direction pulse signals from the ESP32 microcontroller and generated the appropriate '
        'current waveforms for the motor coils. The driver supported input voltages from 20V to 50V DC '
        'and output currents up to 4.2A, configurable via DIP switches. Microstepping was set to 8× '
        '(1600 steps per revolution) to achieve smoother motor operation and reduced vibration. The driver '
        'featured built-in overcurrent, overvoltage, and overtemperature protection. Control signals '
        'included PUL+ (step pulse), DIR+ (direction), and ENA+ (enable), with corresponding negative '
        'terminals connected to ground.'
    )

    # Auger
    p = doc.add_paragraph()
    run = p.add_run('20mm Wood Drill Auger: ')
    run.bold = True
    p.add_run(
        'A 20mm diameter stainless steel wood drill auger bit was adapted as the feed dispensing mechanism. '
        'The auger featured a helical flighting with 20mm pitch, meaning each complete revolution of the '
        'auger advanced the feed material by approximately 20mm along the dispensing tube. Through calibration, '
        'it was determined that one revolution of the auger dispensed approximately 25 grams of commercial '
        'catfish pellets. This volumetric dispensing approach provided consistent feed delivery regardless '
        'of hopper fill level. The auger was housed in a PVC tube connected to the feed hopper, with the '
        'motor shaft coupled directly to the auger via a flexible coupling to accommodate minor misalignment.'
    )

    # HX711 + Load Cell
    p = doc.add_paragraph()
    run = p.add_run('HX711 Load Cell Amplifier and 20kg Load Cell: ')
    run.bold = True
    p.add_run(
        'The HX711 24-bit analog-to-digital converter (ADC) was paired with a 20kg capacity strain gauge '
        'load cell for precise weight measurement of the feed hopper. The HX711 provided high-resolution '
        'weight readings with a sensitivity of approximately 0.1 grams after calibration. The load cell '
        'was mounted beneath the feed hopper to continuously monitor feed level and detect when refilling '
        'was required. The system generated low-feed alerts when the hopper weight dropped below 10% of '
        'capacity (1.5kg remaining from 15kg total capacity). The HX711 communicated with the ESP32 via '
        'a two-wire interface (DOUT and SCK pins), with 10 samples averaged for each reading to reduce noise.'
    )

    # DS18B20
    p = doc.add_paragraph()
    run = p.add_run('DS18B20 Waterproof Temperature Sensor: ')
    run.bold = True
    p.add_run(
        'The DS18B20 digital temperature sensor was used to monitor water temperature in real-time. '
        'This sensor featured a waterproof stainless steel probe with 1-meter cable length, suitable '
        'for submersion in the fish tank. The sensor provided 12-bit resolution (0.0625°C precision) '
        'with an accuracy of ±0.5°C over the range of -10°C to +85°C. Temperature readings were taken '
        'every 5 minutes and used by the Q10 algorithm to adjust feeding rates based on fish metabolic '
        'activity. The sensor communicated via the OneWire protocol, requiring only a single data pin '
        'plus a 4.7kΩ pull-up resistor to 3.3V. Water temperature data were logged to the cloud database '
        'for historical analysis and correlation with growth performance.'
    )

    # JSN-SR04T
    p = doc.add_paragraph()
    run = p.add_run('JSN-SR04T Waterproof Ultrasonic Sensor: ')
    run.bold = True
    p.add_run(
        'The JSN-SR04T waterproof ultrasonic distance sensor was employed for water level monitoring. '
        'This sensor featured an IP67-rated transducer probe suitable for outdoor and wet environments. '
        'The sensor operated by emitting ultrasonic pulses at 40kHz and measuring the time-of-flight '
        'for the echo return, with an effective range of 25cm to 400cm and accuracy of ±1cm. The sensor '
        'was mounted above the water surface to measure the distance to the water level, enabling '
        'detection of water level changes that could affect feeding operations. The sensor required '
        'two GPIO pins: TRIG (trigger pulse output) and ECHO (echo pulse input).'
    )

    # Solar Panel
    p = doc.add_paragraph()
    run = p.add_run('Solar Panel (50-100W Monocrystalline): ')
    run.bold = True
    p.add_run(
        'A monocrystalline solar panel rated at 50-100W served as the primary power source for the '
        'smart feeder system. Monocrystalline panels were selected for their higher efficiency '
        '(typically 18-22%) compared to polycrystalline alternatives, ensuring adequate power generation '
        'even during cloudy conditions. The panel output voltage (typically 18-21V open circuit) was '
        'regulated by the charge controller before charging the 12V battery. The solar-powered design '
        'eliminated the need for grid electricity, making the system suitable for deployment at remote '
        'pond locations. Panel orientation was optimized for maximum solar exposure based on the '
        'installation latitude.'
    )

    # Charge Controller
    p = doc.add_paragraph()
    run = p.add_run('MPPT/PWM Charge Controller: ')
    run.bold = True
    p.add_run(
        'A charge controller (MPPT or PWM type, 10-20A capacity) was used to regulate the charging '
        'of the 12V battery from the solar panel. The controller prevented overcharging and deep '
        'discharge of the battery, extending battery lifespan. MPPT (Maximum Power Point Tracking) '
        'controllers were preferred for their higher efficiency in converting solar panel output to '
        'battery charging current, particularly when panel voltage exceeded battery voltage significantly. '
        'The controller featured LED indicators for charging status and battery level, and provided '
        'load output terminals for powering the feeder electronics.'
    )

    # 12V Battery
    p = doc.add_paragraph()
    run = p.add_run('12V Battery (Lead-Acid or LiFePO4): ')
    run.bold = True
    p.add_run(
        'A 12V rechargeable battery with 12-20Ah capacity provided energy storage for the system, '
        'enabling operation during nighttime and cloudy periods when solar generation was insufficient. '
        'Two battery chemistry options were considered: sealed lead-acid (SLA) batteries offered lower '
        'initial cost but heavier weight and shorter cycle life (300-500 cycles), while lithium iron '
        'phosphate (LiFePO4) batteries provided lighter weight, longer cycle life (2000+ cycles), and '
        'better performance in high-temperature environments typical of tropical fish farming. The '
        'battery capacity was sized to provide at least 24 hours of autonomous operation without solar '
        'charging, accounting for the power consumption of the ESP32, sensors, and periodic motor operation.'
    )

    # LM2596 DC-DC
    p = doc.add_paragraph()
    run = p.add_run('LM2596 DC-DC Step-Down Converter: ')
    run.bold = True
    p.add_run(
        'Two LM2596-based adjustable DC-DC step-down (buck) converter modules were used to regulate '
        'the 12V battery voltage to the required levels for different components. One converter was '
        'set to output 5V for powering the ESP32-CAM and other 5V peripherals, while the second was '
        'set to 3.3V for sensors requiring lower voltage. The LM2596 provided up to 3A output current '
        'with efficiency typically exceeding 90%, minimizing power losses in the voltage conversion '
        'process. The adjustable output voltage was set using the onboard potentiometer and verified '
        'with a multimeter before connecting to sensitive electronics.'
    )

    add_image_placeholder(doc, 'Hardware Block Diagram', '3.2')

    # 3.2.3 Sensor Specifications Table
    add_heading(doc, '3.2.3 Sensor Specifications', level=3)
    sensor_headers = ['Sensor', 'Model', 'Range', 'Accuracy', 'Interface', 'Purpose']
    sensor_data = [
        ['Temperature', 'DS18B20', '-55 to 125°C', '±0.5°C', 'OneWire', 'Water temperature monitoring'],
        ['Weight', 'HX711 + Load Cell', '0-20kg', '±0.1g', 'Digital SPI', 'Feed level measurement'],
        ['Water Level', 'JSN-SR04T', '25-400cm', '±1cm', 'Ultrasonic', 'Water level detection'],
        ['Camera', 'OV2640', '2MP JPEG', 'VGA/SVGA', 'UART', 'Feeding image capture'],
    ]
    add_table(doc, sensor_headers, sensor_data, 'Sensor Specifications', '3.2')

    # =========================================================================
    # 3.2.4 Motor Control System
    # =========================================================================
    add_heading(doc, '3.2.4 Motor Control System', level=3)
    doc.add_paragraph(
        'The feed dispensing mechanism utilized a NEMA 23 bipolar stepper motor coupled with a '
        'DM542 stepper driver and a 20mm wood drill auger. The motor control parameters were '
        'configured to achieve precise feed dispensing with the following specifications:'
    )

    motor_headers = ['Parameter', 'Value', 'Unit', 'Description']
    motor_data = [
        ['Motor Type', 'NEMA 23', '-', 'Bipolar stepper motor'],
        ['Steps per Revolution', '200', 'steps', '1.8 degrees per step'],
        ['Microstepping', '8', 'microsteps', 'Set via DM542 DIP switches'],
        ['Effective Steps/Rev', '1600', 'steps', '200 × 8 microsteps'],
        ['Motor Current', '2.8', 'Amperes', 'Peak current per phase'],
        ['Pulse Width', '5', 'μs', 'Minimum HIGH time for step pulse'],
        ['Step Delay', '1250', 'μs', 'Time between steps at 800 steps/sec'],
        ['Max Speed', '800', 'steps/second', 'Limited for torque at load'],
        ['Motor RPM', '30', 'RPM', '(800 × 60) / 1600 steps'],
        ['Auger Diameter', '20', 'mm', 'Wood drill auger bit'],
        ['Grams per Revolution', '25', 'grams', 'Calibrated feed output'],
        ['Dispensing Rate', '750', 'grams/minute', '30 RPM × 25g per rev'],
    ]
    add_table(doc, motor_headers, motor_data, 'Motor Control Parameters', '3.3')

    add_image_placeholder(doc, 'Motor Control Signal Timing Diagram', '3.3')

    p = doc.add_paragraph()
    run = p.add_run('DM542 DIP Switch Configuration:')
    run.bold = True
    doc.add_paragraph('For 8 microstepping (1600 steps/rev): SW1=OFF, SW2=ON, SW3=ON')
    doc.add_paragraph('For 2.8A motor current: SW4=ON, SW5=OFF, SW6=ON')

    # =========================================================================
    # 3.2.5 Software Technologies
    # =========================================================================
    add_heading(doc, '3.2.5 Software Technologies', level=3)
    doc.add_paragraph(
        'The smart fish feeder system comprised three software layers: embedded firmware running on '
        'the ESP32 microcontroller, a cloud backend providing API services and data storage, and a '
        'mobile application for user interaction and monitoring.'
    )

    software_headers = ['Layer', 'Technology', 'Version', 'Purpose']
    software_data = [
        ['Firmware', 'PlatformIO', '6.x', 'ESP32 development environment'],
        ['Firmware', 'Arduino Framework', '2.x', 'Hardware abstraction layer'],
        ['Firmware', 'TinyGSM', '0.11.x', 'LTE modem communication'],
        ['Backend', 'Go', '1.21+', 'Server-side programming language'],
        ['Backend', 'Gin', '1.9+', 'HTTP web framework'],
        ['Backend', 'PostgreSQL', '15+', 'Relational database'],
        ['Backend', 'Redis', '7+', 'Caching and sessions'],
        ['Mobile', 'Flutter', '3.16+', 'Cross-platform UI framework'],
        ['Mobile', 'Dart', '3.2+', 'Programming language'],
        ['Mobile', 'Riverpod', '2.4+', 'State management'],
        ['Cloud', 'Railway', '-', 'Backend hosting platform'],
        ['Cloud', 'Cloudinary', '-', 'Image storage CDN'],
    ]
    add_table(doc, software_headers, software_data, 'Software Technologies', '3.4')

    # Software descriptions
    p = doc.add_paragraph()
    run = p.add_run('Firmware Layer: ')
    run.bold = True
    p.add_run(
        'The embedded firmware was developed using PlatformIO IDE with the Arduino framework for ESP32. '
        'The firmware implemented sensor reading, motor control, feeding scheduling, Q10 algorithm '
        'calculations, and MQTT communication. The TinyGSM library provided abstraction for the A7670G '
        'LTE modem, enabling TCP/IP connectivity for MQTT messaging. Firmware was organized into modular '
        'components including DeviceManager, SensorManager, FeedingController, CommunicationManager, '
        'and PowerManager for maintainability and testing.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Backend Layer: ')
    run.bold = True
    p.add_run(
        'The cloud backend was developed in Go programming language using the Gin web framework. '
        'The backend provided RESTful API endpoints for user authentication, device management, '
        'feeding control, and data retrieval. PostgreSQL database stored user accounts, device '
        'configurations, feeding schedules, sensor readings, and feeding event logs. Redis provided '
        'caching for frequently accessed data and session management. The backend also implemented '
        'an MQTT client for bidirectional communication with feeder devices, enabling real-time '
        'commands and telemetry. The backend was deployed on Railway cloud platform for scalability '
        'and reliability.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Mobile Application Layer: ')
    run.bold = True
    p.add_run(
        'The mobile application was developed using Flutter framework with Dart programming language, '
        'enabling deployment on both iOS and Android platforms from a single codebase. The application '
        'provided user interface for device pairing, feeding schedule configuration, manual feeding '
        'triggers, real-time sensor monitoring, feeding history visualization, and alert notifications. '
        'Riverpod state management library was used for reactive UI updates and separation of business '
        'logic from presentation. The application communicated with the backend via HTTPS REST API calls '
        'and received push notifications for feeding events and alerts.'
    )

    add_image_placeholder(doc, 'Software Architecture Diagram', '3.4')

    # =========================================================================
    # 3.3 METHODS
    # =========================================================================
    add_heading(doc, '3.3 Methods', level=2)

    add_heading(doc, '3.3.1 System Architecture', level=3)
    doc.add_paragraph(
        'The smart fish feeder system architecture consisted of three main layers: '
        'the hardware layer (sensors and actuators at the fish pond), the cloud infrastructure '
        '(backend services, database, and MQTT broker), and the mobile application layer '
        '(user interface for monitoring and control). Communication between the hardware layer '
        'and cloud infrastructure was established via LTE cellular network using MQTT protocol, '
        'while the mobile application communicated with the backend via HTTPS REST API.'
    )
    add_image_placeholder(doc, 'System Architecture Diagram', '3.5')

    # =========================================================================
    # 3.3.2 Q10 Temperature Algorithm
    # =========================================================================
    add_heading(doc, '3.3.2 Q10 Temperature-Adjusted Feeding Algorithm', level=3)
    doc.add_paragraph(
        'The Q10 model was employed to adjust feeding rates based on water temperature. '
        'The Q10 coefficient represented the factor by which metabolic rate increased for '
        'every 10°C rise in temperature. For African catfish (Clarias gariepinus), a Q10 '
        'value of 2.1 was used based on established literature (Hogendoorn, 1983; Hecht, 2013). '
        'This temperature-dependent feeding approach aimed to match feed supply with the '
        'metabolic demands of the fish, potentially improving feed conversion efficiency '
        'and reducing feed waste.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Q10 Metabolic Adjustment Formula:')
    run.bold = True
    add_formula(doc, 'FR_adj = SFR × Q10^((T - T_ref) / 10) × TIF')

    doc.add_paragraph('Where:')
    doc.add_paragraph('• FR_adj = Adjusted feeding rate (g/day)')
    doc.add_paragraph('• SFR = Standard feeding rate (g/day)')
    doc.add_paragraph('• Q10 = Temperature coefficient (2.1 for African catfish)')
    doc.add_paragraph('• T = Current water temperature (°C)')
    doc.add_paragraph('• T_ref = Reference temperature (25°C)')
    doc.add_paragraph('• TIF = Thermal Inhibition Factor (penalty when T > 32°C or T < 18°C)')

    q10_headers = ['Parameter', 'Symbol', 'Value', 'Unit', 'Description']
    q10_data = [
        ['Q10 Coefficient', 'Q10', '2.1', '-', 'Metabolic rate temperature sensitivity'],
        ['Reference Temperature', 'T_ref', '25', '°C', 'Standard temperature for calculations'],
        ['Optimal Range', 'T_opt', '26-32', '°C', 'Best growth temperature for catfish'],
        ['Critical Low', 'T_crit_low', '18', '°C', 'Feeding reduced below this temperature'],
        ['Critical High', 'T_crit_high', '32', '°C', 'Feeding reduced above this temperature'],
        ['Thermal Inhibition', 'TIF', '0.3-0.8', '-', 'Penalty during thermal stress'],
        ['Standard Feeding Rate', 'SFR', '2-5', '% BW/day', 'Base feeding rate by body weight'],
    ]
    add_table(doc, q10_headers, q10_data, 'Q10 Algorithm Parameters', '3.5')
    add_image_placeholder(doc, 'Feeding Control Algorithm Flowchart', '3.6')

    # =========================================================================
    # 3.3.3 Feed Dispensing Process
    # =========================================================================
    add_heading(doc, '3.3.3 Feed Dispensing Process', level=3)
    doc.add_paragraph(
        'The feed dispensing process involved calculating the required motor steps based on '
        'the target feed amount, executing the motor movement, and verifying the dispensed '
        'quantity using the load cell sensor. The process was initiated either by scheduled '
        'feeding times or manual commands from the mobile application.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Motor Steps Calculation:')
    run.bold = True
    add_formula(doc, 'Total_Steps = (Target_Grams / Grams_Per_Rev) × Steps_Per_Rev × Microsteps')

    p = doc.add_paragraph()
    run = p.add_run('Motor RPM Calculation:')
    run.bold = True
    add_formula(doc, 'RPM = (Steps_Per_Second × 60) / Effective_Steps_Per_Revolution')
    add_formula(doc, 'RPM = (800 × 60) / 1600 = 30 RPM')

    p = doc.add_paragraph()
    run = p.add_run('Dispensing Time Calculation:')
    run.bold = True
    add_formula(doc, 'Time_Seconds = Total_Steps / Steps_Per_Second')

    p = doc.add_paragraph()
    run = p.add_run('Feed Dispensing Rate:')
    run.bold = True
    add_formula(doc, 'Dispensing_Rate = RPM × Grams_Per_Revolution = 30 × 25 = 750 grams/minute')

    calc_headers = ['Step', 'Calculation', 'Value', 'Unit']
    calc_data = [
        ['Target Feed Amount', 'User input', '100', 'grams'],
        ['Q10 Adjustment Factor', 'Q10^((T-25)/10)', '1.15', '-'],
        ['Adjusted Feed Amount', '100 × 1.15', '115', 'grams'],
        ['Revolutions Needed', '115 / 25', '4.6', 'revolutions'],
        ['Steps per Revolution', '200 × 8', '1600', 'steps'],
        ['Total Steps Required', '4.6 × 1600', '7360', 'steps'],
        ['Time per Step', '1 / 800', '1.25', 'milliseconds'],
        ['Total Dispensing Time', '7360 × 1.25', '9.2', 'seconds'],
    ]
    add_table(doc, calc_headers, calc_data, 'Feed Dispensing Calculation Example', '3.6')
    add_image_placeholder(doc, 'Feed Dispensing Mechanism Flowchart', '3.7')

    # =========================================================================
    # 3.3.4 Experimental Design
    # =========================================================================
    add_heading(doc, '3.3.4 Experimental Design', level=3)
    doc.add_paragraph(
        'A completely randomized design (CRD) was employed with two treatment groups. '
        'African catfish fingerlings were randomly assigned to either the control group '
        '(manual feeding) or the treatment group (smart automated feeding). The experiment '
        'was conducted over an 8-week period with weekly measurements of growth parameters.'
    )

    exp_headers = ['Parameter', 'Control Group', 'Treatment Group']
    exp_data = [
        ['Feeding Method', 'Manual', 'Smart Automatic Feeder'],
        ['Feeding Schedule', 'Fixed 3× daily', 'Q10-adjusted dynamic'],
        ['Feed Amount', '3% body weight fixed', 'Temperature-adjusted'],
        ['Temperature Monitoring', 'Manual thermometer', 'Continuous DS18B20 sensor'],
        ['Data Recording', 'Manual logbook', 'Automatic cloud logging'],
        ['Feeding Verification', 'Visual observation', 'ESP32-CAM image capture'],
        ['Number of Fish', '50', '50'],
        ['Tank Size', '500 liters', '500 liters'],
        ['Duration', '8 weeks', '8 weeks'],
        ['Feed Type', 'Commercial catfish pellets', 'Commercial catfish pellets'],
        ['Water Exchange', '30% weekly', '30% weekly'],
    ]
    add_table(doc, exp_headers, exp_data, 'Experimental Design Parameters', '3.7')
    add_image_placeholder(doc, 'Experimental Design Flowchart', '3.8')

    # =========================================================================
    # 3.3.5 Growth Performance Parameters
    # =========================================================================
    add_heading(doc, '3.3.5 Growth Performance Parameters', level=3)
    doc.add_paragraph(
        'The following growth performance parameters were measured and calculated to evaluate '
        'the effect of the smart feeding system on African catfish growth:'
    )

    growth_headers = ['Parameter', 'Formula', 'Unit', 'Description']
    growth_data = [
        ['Weight Gain (WG)', 'WG = W₂ - W₁', 'grams', 'Final weight minus initial weight'],
        ['Specific Growth Rate (SGR)', 'SGR = (ln W₂ - ln W₁) / t × 100', '%/day', 'Daily percentage growth rate'],
        ['Average Daily Gain (ADG)', 'ADG = WG / Days', 'g/day', 'Average weight gained per day'],
        ['Feed Conversion Ratio (FCR)', 'FCR = Feed Consumed / Weight Gain', 'ratio', 'Feed efficiency measure'],
        ['Protein Efficiency Ratio (PER)', 'PER = WG / Protein Intake', 'ratio', 'Protein utilization efficiency'],
        ['Feed Efficiency (FE)', 'FE = (WG / Feed) × 100', '%', 'Percentage of feed converted'],
        ['Condition Factor (K)', 'K = (W / L³) × 100', '-', 'Fish body condition index'],
        ['Survival Rate (SR)', 'SR = (Final / Initial) × 100', '%', 'Percentage of fish surviving'],
    ]
    add_table(doc, growth_headers, growth_data, 'Growth Performance Parameters', '3.8')

    p = doc.add_paragraph()
    run = p.add_run('Specific Growth Rate (SGR) Formula:')
    run.bold = True
    add_formula(doc, 'SGR = ((ln W₂ - ln W₁) / t) × 100')
    doc.add_paragraph('Where:')
    doc.add_paragraph('• W₂ = Final body weight (g)')
    doc.add_paragraph('• W₁ = Initial body weight (g)')
    doc.add_paragraph('• t = Time period (days)')
    doc.add_paragraph('• ln = Natural logarithm')

    p = doc.add_paragraph()
    run = p.add_run('Feed Conversion Ratio (FCR) Formula:')
    run.bold = True
    add_formula(doc, 'FCR = Total Feed Consumed (g) / Total Weight Gain (g)')
    doc.add_paragraph('Note: Lower FCR indicated better feed efficiency. Target FCR for catfish: 1.0 - 1.5')

    # =========================================================================
    # 3.3.6 Statistical Analysis
    # =========================================================================
    add_heading(doc, '3.3.6 Statistical Analysis', level=3)
    doc.add_paragraph(
        'Data collected from the experiment were analyzed using descriptive and inferential '
        'statistics. The Independent Samples t-Test was employed to compare the means of '
        'growth performance parameters between the control group (manual feeding) and the '
        'treatment group (smart automated feeding).'
    )

    p = doc.add_paragraph()
    run = p.add_run('Independent Samples t-Test:')
    run.bold = True
    doc.add_paragraph(
        'The t-Test was chosen as the appropriate statistical method because the study '
        'involved comparing the means of two independent groups (control vs. treatment). '
        'The following assumptions were verified before conducting the t-Test:'
    )
    doc.add_paragraph('1. Independence of observations: Fish were randomly assigned to groups')
    doc.add_paragraph('2. Normal distribution: Data were tested for normality using Shapiro-Wilk test')
    doc.add_paragraph('3. Homogeneity of variance: Levene\'s test was used to verify equal variances')

    p = doc.add_paragraph()
    run = p.add_run('t-Test Formula:')
    run.bold = True
    add_formula(doc, 't = (X̄₁ - X̄₂) / √(s²p × (1/n₁ + 1/n₂))')
    doc.add_paragraph('Where:')
    doc.add_paragraph('• X̄₁, X̄₂ = Sample means of group 1 and group 2')
    doc.add_paragraph('• s²p = Pooled variance')
    doc.add_paragraph('• n₁, n₂ = Sample sizes of group 1 and group 2')

    p = doc.add_paragraph()
    run = p.add_run('Pooled Variance Formula:')
    run.bold = True
    add_formula(doc, 's²p = ((n₁-1)s₁² + (n₂-1)s₂²) / (n₁ + n₂ - 2)')

    p = doc.add_paragraph()
    run = p.add_run('Degrees of Freedom:')
    run.bold = True
    add_formula(doc, 'df = n₁ + n₂ - 2')

    p = doc.add_paragraph()
    run = p.add_run('Statistical Significance:')
    run.bold = True
    doc.add_paragraph(
        'A significance level of α = 0.05 was used for all statistical tests. Results were '
        'considered statistically significant when p < 0.05. The null hypothesis (H₀) stated '
        'that there was no significant difference between the means of the two groups, while '
        'the alternative hypothesis (H₁) stated that there was a significant difference.'
    )
    doc.add_paragraph('• H₀: μ₁ = μ₂ (No significant difference between groups)')
    doc.add_paragraph('• H₁: μ₁ ≠ μ₂ (Significant difference exists between groups)')

    p = doc.add_paragraph()
    run = p.add_run('Decision Rule:')
    run.bold = True
    doc.add_paragraph('• If |t_calculated| > t_critical or p-value < 0.05: Reject H₀')
    doc.add_paragraph('• If |t_calculated| ≤ t_critical or p-value ≥ 0.05: Fail to reject H₀')

    stat_headers = ['Parameter', 'Description']
    stat_data = [
        ['Statistical Test', 'Independent Samples t-Test'],
        ['Significance Level (α)', '0.05'],
        ['Confidence Interval', '95%'],
        ['Software', 'SPSS / Microsoft Excel / Python (SciPy)'],
        ['Normality Test', 'Shapiro-Wilk Test'],
        ['Variance Test', 'Levene\'s Test'],
        ['Effect Size', 'Cohen\'s d'],
    ]
    add_table(doc, stat_headers, stat_data, 'Statistical Analysis Methods', '3.9')

    # =========================================================================
    # 3.3.7 Hardware Wiring Connections
    # =========================================================================
    add_heading(doc, '3.3.7 Hardware Wiring Connections', level=3)
    doc.add_paragraph(
        'The following wiring connections were established between the microcontroller '
        'and peripheral components:'
    )

    p = doc.add_paragraph()
    run = p.add_run('Motor Driver Connections (DM542):')
    run.bold = True
    motor_wire_headers = ['ESP32 Pin', 'DM542 Pin', 'Function']
    motor_wire_data = [
        ['GPIO32', 'PUL+', 'Step pulse signal'],
        ['GPIO33', 'DIR+', 'Direction control'],
        ['GPIO0', 'ENA+', 'Enable (active LOW)'],
        ['GND', 'PUL-, DIR-, ENA-', 'Common ground'],
        ['-', 'Power (12-48V DC)', 'Motor power supply'],
        ['-', 'A+, A-, B+, B-', 'NEMA 23 motor coils'],
    ]
    add_table(doc, motor_wire_headers, motor_wire_data, 'Motor Driver Wiring', '3.10')

    p = doc.add_paragraph()
    run = p.add_run('Sensor Connections:')
    run.bold = True
    sensor_wire_headers = ['Sensor', 'ESP32 Pin', 'Function']
    sensor_wire_data = [
        ['HX711 DOUT', 'GPIO39 (VN)', 'Load cell data output'],
        ['HX711 SCK', 'GPIO5', 'Load cell clock'],
        ['DS18B20 Data', 'GPIO23', 'Temperature data (with 4.7kΩ pullup)'],
        ['JSN-SR04T TRIG', 'GPIO17', 'Ultrasonic trigger'],
        ['JSN-SR04T ECHO', 'GPIO34', 'Ultrasonic echo'],
    ]
    add_table(doc, sensor_wire_headers, sensor_wire_data, 'Sensor Wiring Connections', '3.11')
    add_image_placeholder(doc, 'Complete Wiring Diagram', '3.9')

    # =========================================================================
    # 3.4 DATA COLLECTION
    # =========================================================================
    add_heading(doc, '3.4 Data Collection', level=2)
    doc.add_paragraph(
        'Data collection was performed automatically by the smart feeder system and manually '
        'for certain parameters. The following data were collected throughout the experiment:'
    )
    doc.add_paragraph('• Water temperature: Recorded every 5 minutes via DS18B20 sensor')
    doc.add_paragraph('• Feed dispensed: Recorded for each feeding event via HX711 load cell')
    doc.add_paragraph('• Feeding images: Captured via ESP32-CAM for each feeding event')
    doc.add_paragraph('• Fish weight: Measured weekly using digital scale (individual sampling)')
    doc.add_paragraph('• Fish length: Measured weekly using measuring board')
    doc.add_paragraph('• Mortality: Recorded daily')
    doc.add_paragraph('• Water quality: Monitored weekly (pH, ammonia, nitrite)')

    doc.add_paragraph(
        'All sensor data were transmitted via LTE cellular connection to the cloud backend '
        'and stored in a PostgreSQL database for analysis. The mobile application provided '
        'real-time monitoring and historical data visualization. Feeding images were stored '
        'on Cloudinary CDN and linked to feeding event records in the database.'
    )

    # =========================================================================
    # 3.5 ETHICAL CONSIDERATIONS
    # =========================================================================
    add_heading(doc, '3.5 Ethical Considerations', level=2)
    doc.add_paragraph(
        'This study was conducted in accordance with ethical guidelines for animal research. '
        'The fish were handled humanely, and proper care was taken to minimize stress during '
        'measurements. Water quality was maintained within optimal ranges for African catfish, '
        'and any fish showing signs of disease or distress were promptly treated or removed '
        'from the experiment. Feeding rates were designed to meet nutritional requirements '
        'without overfeeding, which could lead to water quality deterioration and health issues.'
    )

    # =========================================================================
    # Save Document
    # =========================================================================
    doc.save('docs/CHAPTER_3_PROPOSAL.docx')
    print('Document saved: docs/CHAPTER_3_PROPOSAL.docx')

if __name__ == '__main__':
    create_chapter3_document()

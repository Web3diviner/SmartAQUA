"""
Script to read and extract text from the Darasimi proposal.docx file
"""

from docx import Document

def read_docx(filepath):
    """Read a docx file and extract all text content"""
    doc = Document(filepath)
    
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract tables
    for table_idx, table in enumerate(doc.tables):
        full_text.append(f"\n--- TABLE {table_idx + 1} ---")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            full_text.append(" | ".join(row_text))
    
    return "\n".join(full_text)

if __name__ == "__main__":
    filepath = "Darasimi proposal.docx"
    content = read_docx(filepath)
    
    # Save to text file for review
    with open("docs/proposal_content.txt", "w", encoding="utf-8") as f:
        f.write(content)
    
    print(f"Content extracted and saved to docs/proposal_content.txt")
    print(f"Total characters: {len(content)}")
    print("\n" + "="*50)
    print("FIRST 5000 CHARACTERS:")
    print("="*50)
    print(content[:5000])

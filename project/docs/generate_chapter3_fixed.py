"""
Generate Chapter 3: Materials and Methods - FIXED VERSION
Fixes:
1. Added Nutrient Utilization Parameters (APD, PER, NPU, LR, etc.)
2. Tables with 3 horizontal lines only (APA style)
3. In-text references to all tables and figures
4. More specific section headings
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

def set_table_borders_apa(table):
    """Set APA style borders - only 3 horizontal lines (top, below header, bottom)"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    tblBorders = OxmlElement('w:tblBorders')
    
    # Top border
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)
    
    # Bottom border
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)
    
    # No left, right, insideV borders
    for border_name in ['left', 'right', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    
    # InsideH - only for header row (we'll handle this per row)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'nil')
    tblBorders.append(insideH)
    
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def add_header_row_border(table):
    """Add bottom border to header row only"""
    row = table.rows[0]
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), '000000')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

def add_table_apa(doc, headers, data, caption, table_num):
    """Add APA-style table with 3 lines only"""
    # Table caption (above table, italicized)
    cap_p = doc.add_paragraph()
    run = cap_p.add_run(f'Table {table_num}')
    run.bold = True
    cap_p.add_run(f'\n{caption}')
    cap_p.runs[1].italic = True
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
    
    # Apply APA borders
    set_table_borders_apa(table)
    add_header_row_border(table)
    
    doc.add_paragraph()
    return table

def add_image_placeholder(doc, caption, fig_num):
    """Add figure placeholder with caption below"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n[INSERT FIGURE {fig_num} HERE]\n')
    run.font.size = Pt(11)
    
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(f'Figure {fig_num}. ')
    run.bold = True
    run.italic = True
    cap_p.add_run(caption)
    cap_p.runs[1].italic = True
    doc.add_paragraph()

def add_formula(doc, formula_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula_text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)

def create_chapter3():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # TITLE
    title = doc.add_heading('CHAPTER THREE', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading('MATERIALS AND METHODS', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
